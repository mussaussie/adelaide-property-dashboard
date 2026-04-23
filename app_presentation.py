"""
Adelaide Property Decision Dashboard

A presentation-focused Streamlit app that reuses the existing output datasets
without changing app.py. It adds stronger first-screen storytelling, suburb
comparison, opportunity ranking, metric-switching maps, and richer ML/risk
visuals.
"""

from __future__ import annotations

import copy
import io
import json
import os
from html import escape

import folium
import numpy as np
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import streamlit as st
from docx import Document
from fpdf import FPDF, XPos, YPos
from streamlit_folium import st_folium


BASE_DIR = os.path.dirname(os.path.abspath(__file__))

MASTER_COLS = [
    "Suburb",
    "Avg_Price_All_Time",
    "Median_Price_All_Time",
    "Min_Price_Ever",
    "Max_Price_Ever",
    "Price_Volatility",
    "Quarter_Count",
    "Current_Price_2025",
    "First_Price_2019",
    "Price_Growth_Amount",
    "Price_Growth_Percent",
    "Total_Crime_Count",
    "Property_Crime_Count",
    "Person_Crime_Count",
    "Crime_OFFENCES_AGAINST_PROPERTY",
    "Crime_OFFENCES_AGAINST_THE_PERSON",
    "G01_Population_Total",
    "G02_Median_age_persons",
    "G02_Median_mortgage_repay_monthly",
    "G02_Median_tot_prsnl_inc_weekly",
    "G02_Median_rent_weekly",
    "G02_Median_tot_fam_inc_weekly",
    "G02_Median_tot_hhd_inc_weekly",
    "G02_Average_household_size",
]

RISK_COLS = [
    "Suburb",
    "Actual_Price_2025",
    "Predicted_Price_2025",
    "Prediction_Error",
    "Error_Percent",
    "Absolute_Error",
    "Historical_Growth_Rate",
    "Forecast_Price_2026",
    "Forecast_2026_Lower",
    "Forecast_2026_Upper",
    "Median_Income",
    "Crime_Rate",
    "Education_Rate",
    "Employment_Rate",
    "Expected_Growth_2026",
    "Value_Gap",
    "Value_Gap_Percent",
    "Value_Category",
    "Value_Score",
    "Growth_Score",
    "Safety_Score",
    "Income_Score",
    "Education_Score",
    "Investment_Score",
    "Price_Category",
    "Crime_Rate_Risk",
    "Crime_Risk_Category",
    "Economic_Risk",
    "Economic_Risk_Category",
    "Prediction_Uncertainty_Risk",
    "Growth_Risk",
    "Price_Level_Risk",
    "Market_Risk",
    "Market_Risk_Category",
    "Total_Risk_Score",
    "Total_Risk_Category",
    "Risk_Adjusted_Return",
    "Investment_Strategy",
]

RENTAL_COLS = [
    "Suburb",
    "Census_Rent_2021",
    "Fair_Rent_2025",
    "Fair_House_Rent_2025",
    "Fair_Unit_Rent_2025",
    "Estimated_Actual_Rent_2025",
    "Actual_House_Rent_2025",
    "Actual_Unit_Rent_2025",
    "Greediness_Percent",
    "Fair_House_Yield",
    "Actual_House_Yield",
    "Actual_Unit_Yield",
    "Affordability_Category",
    "Affordability_Ratio",
    "Individual_Affordability",
    "Household_Affordability",
]

CULTURE_COLS = [
    "Suburb",
    "Indian_Population",
    "Indian_Percent",
    "Chinese_Population",
    "Chinese_Percent",
    "Vietnamese_Population",
    "Vietnamese_Percent",
    "Italian_Population",
    "Italian_Percent",
    "Greek_Population",
    "Greek_Percent",
    "Cultural_Diversity_Index",
]

CRIME_COLS = ["Suburb", "No_of_Crimes", "Crime_Type_1", "Crime_Type_2", "Crime_Type_3"]
TS_COLS = ["Suburb", "Median_Price", "Period", "Quarter", "Year"]

METRIC_OPTIONS = {
    "Latest price": ("Current_Price_2025", "Price", "price"),
    "29-quarter growth": ("Price_Growth_Percent", "Growth %", "growth"),
    "Next-year growth": ("Expected_Growth_2026", "Growth %", "growth"),
    "Risk score": ("Total_Risk_Score", "Risk", "risk"),
    "House yield": ("Actual_House_Yield", "Yield %", "yield"),
    "Crime rate": ("Crime_Rate_Per_1000", "Crime / 1k", "risk"),
    "Cultural diversity": ("Cultural_Diversity_Index", "Diversity", "diversity"),
}

DISPLAY_COLUMN_LABELS = {
    "Current_Price_2025": "Latest Price",
    "Actual_Price_2025": "Latest Price",
    "Predicted_Price_2025": "Predicted Latest Price",
    "Forecast_Price_2026": "Next-Year Forecast",
    "Forecast_2026_Lower": "Next-Year Forecast Lower",
    "Forecast_2026_Upper": "Next-Year Forecast Upper",
    "Price_Growth_Percent": "Q1 2019 to Q1 2026 Growth",
    "Expected_Growth_2026": "Next-Year Growth",
    "Actual_House_Yield": "House Yield",
    "Crime_Rate_Per_1000": "Crime / 1k",
    "Total_Risk_Category": "Risk Category",
    "Total_Risk_Score": "Risk Score",
    "Opportunity_Score": "Opportunity Score",
    "Investment_Strategy": "Investment Strategy",
    "G02_Median_tot_hhd_inc_weekly": "Median Household Income / Week",
}

PALETTES = {
    "price": ["#e7f3ec", "#b8dec7", "#78bd96", "#31866a", "#0f4f47"],
    "growth": ["#eef1f4", "#bdd7ea", "#75add1", "#2f7fb7", "#174b7a"],
    "risk": ["#e3f2e3", "#f1daa0", "#eea35f", "#d9654f", "#9f2f3c"],
    "yield": ["#f6efe3", "#ead6a5", "#d2ab59", "#987a2e", "#59491f"],
    "diversity": ["#eef2f1", "#cad8d7", "#8fb7b5", "#4d918d", "#21635f"],
}


st.set_page_config(
    page_title="Adelaide Property Decision Dashboard",
    page_icon="AP",
    layout="wide",
    initial_sidebar_state="collapsed",
)


st.markdown(
    """
<style>
    :root {
        --page: #f5f6f1;
        --panel: #ffffff;
        --ink: #18222f;
        --muted: #647181;
        --line: #d9dfdc;
        --teal: #176b5b;
        --blue: #285f96;
        --coral: #c94e3f;
        --amber: #a87922;
        --soft-teal: #e5f1ed;
        --soft-blue: #e8eef6;
        --soft-coral: #f8e9e5;
        --soft-amber: #f6eddc;
    }

    .stApp {
        background: var(--page);
        color: var(--ink);
    }

    html, body, [class*="css"] {
        font-family: Inter, Segoe UI, Roboto, Arial, sans-serif;
    }

    section[data-testid="stSidebar"] {
        background: #ffffff;
        border-right: 1px solid var(--line);
    }

    .block-container {
        padding-top: 1.4rem;
        padding-bottom: 3rem;
        max-width: 1480px;
    }

    .topbar {
        background: #ffffff;
        border: 1px solid var(--line);
        border-radius: 8px;
        padding: 1.15rem 1.25rem;
        margin-bottom: 1rem;
        box-shadow: 0 12px 30px rgba(24, 34, 47, 0.06);
    }

    .eyebrow {
        color: var(--teal);
        font-size: 0.78rem;
        font-weight: 800;
        text-transform: uppercase;
        letter-spacing: 0;
        margin-bottom: 0.25rem;
    }

    .topbar h1 {
        color: var(--ink);
        font-size: 2.15rem;
        line-height: 1.05;
        margin: 0;
        font-weight: 850;
        letter-spacing: 0;
    }

    .topbar p {
        color: var(--muted);
        margin: 0.45rem 0 0 0;
        font-size: 1rem;
    }

    .section-title {
        color: var(--ink);
        font-size: 1.25rem;
        font-weight: 800;
        margin: 1.25rem 0 0.55rem 0;
    }

    .section-subtitle {
        color: var(--muted);
        font-size: 0.92rem;
        margin: -0.25rem 0 0.75rem 0;
    }

    .metric-card {
        min-height: 118px;
        background: #ffffff;
        border: 1px solid var(--line);
        border-radius: 8px;
        padding: 1rem;
        box-shadow: 0 10px 24px rgba(24, 34, 47, 0.06);
    }

    .metric-card.teal { border-top: 4px solid var(--teal); }
    .metric-card.blue { border-top: 4px solid var(--blue); }
    .metric-card.coral { border-top: 4px solid var(--coral); }
    .metric-card.amber { border-top: 4px solid var(--amber); }

    .metric-label {
        color: var(--muted);
        font-size: 0.74rem;
        font-weight: 800;
        text-transform: uppercase;
        letter-spacing: 0;
        line-height: 1.25;
    }

    .metric-value {
        color: var(--ink);
        font-size: 1.62rem;
        line-height: 1.12;
        font-weight: 850;
        margin-top: 0.35rem;
        overflow-wrap: anywhere;
    }

    .metric-note {
        color: var(--muted);
        font-size: 0.82rem;
        margin-top: 0.35rem;
        line-height: 1.3;
    }

    .insight-box {
        background: #ffffff;
        border: 1px solid var(--line);
        border-left: 5px solid var(--teal);
        border-radius: 8px;
        padding: 1rem 1.1rem;
        box-shadow: 0 10px 24px rgba(24, 34, 47, 0.05);
    }

    .insight-box strong {
        color: var(--ink);
    }

    .pill-row {
        display: flex;
        flex-wrap: wrap;
        gap: 0.45rem;
        margin-top: 0.65rem;
    }

    .pill {
        display: inline-flex;
        align-items: center;
        background: var(--soft-blue);
        border: 1px solid #c8d6e8;
        color: #254864;
        border-radius: 999px;
        padding: 0.32rem 0.65rem;
        font-size: 0.8rem;
        font-weight: 750;
    }

    .pill.teal { background: var(--soft-teal); border-color: #b9d6cd; color: var(--teal); }
    .pill.coral { background: var(--soft-coral); border-color: #efc8bf; color: var(--coral); }
    .pill.amber { background: var(--soft-amber); border-color: #e5cf9d; color: var(--amber); }

    .rank-card {
        background: #ffffff;
        border: 1px solid var(--line);
        border-radius: 8px;
        padding: 0.85rem 0.95rem;
        margin-bottom: 0.65rem;
        box-shadow: 0 8px 20px rgba(24, 34, 47, 0.045);
    }

    .rank-title {
        display: flex;
        align-items: center;
        justify-content: space-between;
        gap: 0.75rem;
    }

    .rank-title strong {
        color: var(--ink);
        font-size: 1.02rem;
    }

    .rank-score {
        color: #ffffff;
        background: var(--teal);
        border-radius: 999px;
        padding: 0.25rem 0.55rem;
        font-size: 0.78rem;
        font-weight: 800;
        white-space: nowrap;
    }

    .rank-meta {
        color: var(--muted);
        font-size: 0.83rem;
        margin-top: 0.35rem;
        line-height: 1.35;
    }

    .note-box {
        color: #55471f;
        background: #fbf5e7;
        border: 1px solid #ead9ad;
        border-radius: 8px;
        padding: 0.9rem 1rem;
        font-size: 0.9rem;
        line-height: 1.45;
    }

    .action-button button {
        min-height: 3.1rem;
    }

    .stButton > button {
        background: linear-gradient(135deg, #176b5b, #285f96) !important;
        color: #ffffff !important;
        border: 0 !important;
        border-radius: 8px !important;
        font-weight: 850 !important;
        box-shadow: 0 10px 22px rgba(24, 34, 47, 0.18) !important;
        min-height: 2.8rem;
    }

    .stButton > button:hover {
        filter: brightness(1.08);
        transform: translateY(-1px);
    }

    .stDownloadButton > button {
        background: linear-gradient(135deg, #a87922, #c94e3f) !important;
        color: #ffffff !important;
        border: 0 !important;
        border-radius: 8px !important;
        font-weight: 850 !important;
        box-shadow: 0 10px 22px rgba(24, 34, 47, 0.16) !important;
    }

    .map-control-title {
        color: var(--ink);
        font-size: 1.35rem;
        font-weight: 850;
        margin: 0.4rem 0 0.25rem 0;
    }

    .map-control-help {
        color: var(--muted);
        font-size: 0.92rem;
        margin-bottom: 0.35rem;
    }

    .glossary-term {
        background: #ffffff;
        border: 1px solid var(--line);
        border-radius: 8px;
        padding: 0.9rem 1rem;
        margin-bottom: 0.65rem;
        box-shadow: 0 8px 20px rgba(24, 34, 47, 0.045);
    }

    .glossary-term strong {
        color: var(--ink);
        display: block;
        font-size: 1rem;
        margin-bottom: 0.25rem;
    }

    div[data-testid="stDataFrame"] {
        border: 1px solid var(--line);
        border-radius: 8px;
        overflow: hidden;
    }

    .stTabs [data-baseweb="tab-list"] {
        gap: 0.35rem;
        border-bottom: 1px solid var(--line);
    }

    .stTabs [data-baseweb="tab"] {
        border-radius: 8px 8px 0 0;
        padding: 0.5rem 0.8rem;
        color: var(--muted);
        font-weight: 750;
    }

    .stTabs [aria-selected="true"] {
        color: var(--teal) !important;
        background: #ffffff !important;
        border: 1px solid var(--line);
        border-bottom-color: #ffffff;
    }

    @media screen and (max-width: 900px) {
        .block-container {
            padding-left: 0.85rem;
            padding-right: 0.85rem;
            padding-top: 0.75rem;
        }

        .topbar {
            padding: 0.95rem;
            margin-bottom: 0.75rem;
        }

        .topbar h1 {
            font-size: 1.45rem;
            line-height: 1.12;
        }

        .topbar p {
            font-size: 0.9rem;
            line-height: 1.35;
        }

        .section-title {
            font-size: 1.08rem;
            margin-top: 1rem;
        }

        .section-subtitle {
            font-size: 0.84rem;
            line-height: 1.35;
        }

        .metric-card {
            min-height: auto;
            padding: 0.85rem;
            margin-bottom: 0.35rem;
        }

        .metric-label {
            font-size: 0.68rem;
        }

        .metric-value {
            font-size: 1.18rem;
        }

        .metric-note,
        .rank-meta,
        .note-box,
        .insight-box {
            font-size: 0.82rem;
        }

        .rank-card {
            padding: 0.8rem;
            margin-bottom: 0.5rem;
        }

        .rank-title {
            align-items: flex-start;
        }

        .rank-title strong {
            font-size: 0.92rem;
        }

        .pill {
            font-size: 0.72rem;
            padding: 0.26rem 0.5rem;
        }

        .map-control-title {
            font-size: 1.12rem;
        }

        .stTabs [data-baseweb="tab-list"] {
            flex-wrap: wrap;
            gap: 0.25rem;
        }

        .stTabs [data-baseweb="tab"] {
            font-size: 0.78rem;
            padding: 0.4rem 0.55rem;
        }

        iframe {
            min-height: 420px;
        }
    }

    @media screen and (max-width: 520px) {
        .block-container {
            padding-left: 0.65rem;
            padding-right: 0.65rem;
        }

        .topbar h1 {
            font-size: 1.28rem;
        }

        .eyebrow {
            font-size: 0.68rem;
        }

        .metric-value {
            font-size: 1.05rem;
        }

        .rank-score {
            font-size: 0.7rem;
            padding: 0.2rem 0.45rem;
        }
    }
</style>
""",
    unsafe_allow_html=True,
)


def read_csv(path: str, usecols: list[str] | None = None) -> pd.DataFrame:
    full_path = os.path.join(BASE_DIR, path)
    if not os.path.exists(full_path):
        return pd.DataFrame()
    if usecols is None:
        return pd.read_csv(full_path)
    header = pd.read_csv(full_path, nrows=0).columns.tolist()
    valid_cols = [col for col in usecols if col in header]
    if not valid_cols:
        return pd.DataFrame()
    return pd.read_csv(full_path, usecols=valid_cols)


def merge_new_columns(base: pd.DataFrame, other: pd.DataFrame) -> pd.DataFrame:
    if other.empty or "Suburb" not in other.columns:
        return base
    cols = ["Suburb"] + [col for col in other.columns if col != "Suburb" and col not in base.columns]
    if len(cols) <= 1:
        return base
    return base.merge(other[cols], on="Suburb", how="left")


@st.cache_data
def load_data() -> tuple[pd.DataFrame, pd.DataFrame]:
    master = read_csv("data/clean/master_dataset_by_suburb.csv", MASTER_COLS)
    risk = read_csv("data/risk_analysis/complete_risk_analysis.csv", RISK_COLS)
    rental = read_csv("data/rental/complete_rental_analysis.csv", RENTAL_COLS)
    culture = read_csv("data/demographics/cultural_demographics.csv", CULTURE_COLS)
    crime = read_csv("data/suburb_crime_offense_analysis.csv", CRIME_COLS)
    ts = read_csv("data/clean/property_timeseries_2019_2025.csv", TS_COLS)

    df = master.copy()
    for source in (risk, rental, culture, crime):
        df = merge_new_columns(df, source)

    if "Crime_Rate_Per_1000" not in df.columns:
        if {"Total_Crime_Count", "G01_Population_Total"}.issubset(df.columns):
            df["Crime_Rate_Per_1000"] = np.where(
                df["G01_Population_Total"] > 0,
                df["Total_Crime_Count"] / df["G01_Population_Total"] * 1000,
                np.nan,
            )

    score_parts = []
    for col in ["Investment_Score", "Growth_Score", "Safety_Score", "Value_Score"]:
        if col in df.columns:
            score_parts.append(pd.to_numeric(df[col], errors="coerce"))
    if score_parts:
        base_score = pd.concat(score_parts, axis=1).mean(axis=1)
    else:
        base_score = pd.Series(np.nan, index=df.index)

    forecast_bonus = pd.to_numeric(df.get("Expected_Growth_2026"), errors="coerce").fillna(-20) * 0.35
    yield_bonus = pd.to_numeric(df.get("Actual_House_Yield"), errors="coerce").fillna(0) * 2
    risk_penalty = pd.to_numeric(df.get("Total_Risk_Score"), errors="coerce")
    risk_penalty = risk_penalty.fillna(risk_penalty.median()) * 0.25
    unknown_growth_penalty = (
        df.get("Investment_Strategy", pd.Series("", index=df.index))
        .fillna("")
        .str.contains("Unknown Growth", case=False, regex=False)
        .astype(int)
        * 20
    )
    df["Opportunity_Score"] = base_score + forecast_bonus + yield_bonus - risk_penalty - unknown_growth_penalty

    if not ts.empty:
        ts["Year"] = pd.to_numeric(ts["Year"], errors="coerce")
        ts["Median_Price"] = pd.to_numeric(ts["Median_Price"], errors="coerce")
        ts = ts.dropna(subset=["Suburb", "Year", "Median_Price"]).copy()
        ts["Year"] = ts["Year"].astype(int)

    return df, ts


@st.cache_data
def load_geojson() -> dict | None:
    path = os.path.join(BASE_DIR, "adelaide_suburbs.geojson")
    if not os.path.exists(path):
        return None
    with open(path, "r", encoding="utf-8") as handle:
        return json.load(handle)


@st.cache_data
def load_coordinates() -> dict:
    path = os.path.join(BASE_DIR, "suburb_coordinates.json")
    if not os.path.exists(path):
        return {}
    with open(path, "r", encoding="utf-8") as handle:
        return json.load(handle)


def fmt_price(value: float | int | None) -> str:
    if pd.isna(value):
        return "N/A"
    value = float(value)
    if abs(value) >= 1_000_000:
        return f"${value / 1_000_000:.2f}M"
    return f"${value:,.0f}"


def fmt_price_change(value: float | int | None) -> str:
    if pd.isna(value):
        return "N/A"
    value = float(value)
    sign = "-" if value < 0 else "+"
    abs_value = abs(value)
    if abs_value >= 1_000_000:
        return f"{sign}${abs_value / 1_000_000:.2f}M"
    return f"{sign}${abs_value:,.0f}"


def fmt_money(value: float | int | None) -> str:
    if pd.isna(value):
        return "N/A"
    return f"${float(value):,.0f}"


def fmt_num(value: float | int | None, digits: int = 0) -> str:
    if pd.isna(value):
        return "N/A"
    return f"{float(value):,.{digits}f}"


def fmt_pct(value: float | int | None, digits: int = 1) -> str:
    if pd.isna(value):
        return "N/A"
    return f"{float(value):.{digits}f}%"


def value_gap_note(value: float | int | None, category: object) -> tuple[str, str]:
    if pd.isna(value):
        return "N/A", "Model comparison unavailable"
    pct = float(value)
    category_text = str(category) if category is not None and not pd.isna(category) else "No category"
    if pct < 0:
        direction = f"Actual is {abs(pct):.1f}% above model"
    elif pct > 0:
        direction = f"Actual is {pct:.1f}% below model"
    else:
        direction = "Actual matches model"
    return fmt_pct(pct), f"{direction} | {category_text}"


def risk_plain_language(row: pd.Series) -> str:
    risk_category = str(row.get("Total_Risk_Category", "N/A"))
    strategy = str(row.get("Investment_Strategy", "N/A"))
    if "Very Low" in risk_category or risk_category == "Low Risk":
        verdict = "This suburb looks relatively safe in the model."
    elif "Moderate" in risk_category:
        verdict = "This suburb sits in the middle: not the safest, not the riskiest."
    elif "High" in risk_category:
        verdict = "This suburb needs more caution because the model sees higher risk."
    else:
        verdict = "Risk verdict is limited because some model inputs are missing."
    return f"{verdict} Strategy: {strategy}."


def rent_plain_language(row: pd.Series) -> str:
    fair_rent = row.get("Fair_House_Rent_2025")
    actual_rent = row.get("Actual_House_Rent_2025")
    greed = row.get("Greediness_Percent")
    affordability = row.get("Affordability_Category", "N/A")
    if pd.isna(fair_rent) or pd.isna(actual_rent):
        return "Rental estimate is unavailable for this suburb."
    difference = actual_rent - fair_rent
    if pd.isna(greed):
        gap_text = f"about {fmt_money(difference)} per week above the fair-rent estimate"
    elif greed > 10:
        gap_text = f"about {fmt_pct(greed)} above the fair-rent estimate"
    elif greed < -5:
        gap_text = f"below the fair-rent estimate by about {fmt_pct(abs(greed))}"
    else:
        gap_text = "close to the fair-rent estimate"
    return f"Estimated market rent is {gap_text}. Affordability category: {affordability}."


def safe_text(value: object) -> str:
    if value is None or pd.isna(value):
        return "N/A"
    return escape(str(value))


def plain_text(value: object) -> str:
    if value is None or pd.isna(value):
        return "N/A"
    return str(value).encode("latin-1", "replace").decode("latin-1")


def text_bar(value: float | int | None, max_value: float = 100, width: int = 18) -> str:
    if pd.isna(value) or max_value <= 0:
        return "[no data]"
    filled = int(max(0, min(float(value), max_value)) / max_value * width)
    return "[" + "#" * filled + "-" * (width - filled) + "]"


def card(label: str, value: str, note: str = "", tone: str = "teal") -> str:
    return f"""
    <div class="metric-card {tone}">
        <div class="metric-label">{escape(label)}</div>
        <div class="metric-value">{escape(value)}</div>
        <div class="metric-note">{escape(note)}</div>
    </div>
    """


def title_block(title: str, subtitle: str, eyebrow: str = "Property intelligence") -> None:
    st.markdown(
        f"""
        <div class="topbar">
            <div class="eyebrow">{escape(eyebrow)}</div>
            <h1>{escape(title)}</h1>
            <p>{escape(subtitle)}</p>
        </div>
        """,
        unsafe_allow_html=True,
    )


def section(title: str, subtitle: str = "") -> None:
    st.markdown(f'<div class="section-title">{escape(title)}</div>', unsafe_allow_html=True)
    if subtitle:
        st.markdown(f'<div class="section-subtitle">{escape(subtitle)}</div>', unsafe_allow_html=True)


def apply_chart_style(fig: go.Figure, height: int = 390) -> go.Figure:
    fig.update_layout(
        height=height,
        margin=dict(l=20, r=20, t=45, b=35),
        paper_bgcolor="rgba(0,0,0,0)",
        plot_bgcolor="#ffffff",
        font=dict(family="Inter, Segoe UI, Arial", color="#18222f", size=12),
        title_font=dict(size=16, color="#18222f"),
        legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1),
    )
    fig.update_xaxes(gridcolor="#eef1ef", zerolinecolor="#d9dfdc")
    fig.update_yaxes(gridcolor="#eef1ef", zerolinecolor="#d9dfdc")
    return fig


def make_price_history(ts: pd.DataFrame, suburbs: list[str], forecast_df: pd.DataFrame | None = None) -> go.Figure:
    fig = go.Figure()
    palette = ["#176b5b", "#285f96", "#c94e3f", "#a87922", "#5b6573", "#8a5a2b"]
    for idx, suburb in enumerate(suburbs):
        sub_ts = ts[ts["Suburb"] == suburb].copy()
        if sub_ts.empty:
            continue
        sub_ts = sub_ts.sort_values(["Year", "Quarter", "Period"])
        fig.add_trace(
            go.Scatter(
                x=sub_ts["Period"],
                y=sub_ts["Median_Price"],
                mode="lines+markers",
                name=suburb,
                line=dict(color=palette[idx % len(palette)], width=3),
                marker=dict(size=6),
                hovertemplate="<b>%{fullData.name}</b><br>%{x}<br>%{y:$,.0f}<extra></extra>",
            )
        )

    if forecast_df is not None and len(suburbs) == 1:
        row = forecast_df[forecast_df["Suburb"] == suburbs[0]]
        if not row.empty:
            record = row.iloc[0]
            if not pd.isna(record.get("Forecast_Price_2026")):
                fig.add_trace(
                    go.Scatter(
                        x=["Next-Year Forecast"],
                        y=[record["Forecast_Price_2026"]],
                        mode="markers",
                        name="Next-year forecast",
                        marker=dict(size=13, color="#c94e3f", symbol="diamond"),
                        hovertemplate="Next-year forecast<br>%{y:$,.0f}<extra></extra>",
                    )
                )
                if not pd.isna(record.get("Forecast_2026_Lower")) and not pd.isna(record.get("Forecast_2026_Upper")):
                    fig.add_trace(
                        go.Scatter(
                            x=["Next-Year Forecast", "Next-Year Forecast"],
                            y=[record["Forecast_2026_Lower"], record["Forecast_2026_Upper"]],
                            mode="lines",
                            name="Forecast range",
                            line=dict(color="#c94e3f", width=8),
                            hovertemplate="Forecast range<br>%{y:$,.0f}<extra></extra>",
                        )
                    )

    fig.update_layout(title="Price history and next-year forecast", yaxis_title="Median price", xaxis_title="")
    fig.update_yaxes(tickprefix="$", tickformat=",.0f")
    return apply_chart_style(fig, height=430)


def annual_growth_data(ts: pd.DataFrame, suburb: str) -> pd.DataFrame:
    suburb_ts = ts[ts["Suburb"] == suburb].copy()
    if suburb_ts.empty:
        return pd.DataFrame()
    annual = suburb_ts.groupby("Year", as_index=False)["Median_Price"].median()
    annual = annual.sort_values("Year").reset_index(drop=True)
    annual["YoY_Growth_Pct"] = annual["Median_Price"].pct_change() * 100
    annual["YoY_Growth_Amount"] = annual["Median_Price"].diff()
    annual["YoY_Growth_Pct"] = annual["YoY_Growth_Pct"].replace([np.inf, -np.inf], np.nan)
    return annual


def yoy_growth_chart(ts: pd.DataFrame, suburb: str) -> go.Figure | None:
    annual = annual_growth_data(ts, suburb)
    plot = annual.dropna(subset=["YoY_Growth_Pct"]).copy()
    if plot.empty:
        return None

    best_idx = plot["YoY_Growth_Pct"].idxmax()
    worst_idx = plot["YoY_Growth_Pct"].idxmin()
    bar_colors = []
    for idx, value in plot["YoY_Growth_Pct"].items():
        if idx == best_idx:
            bar_colors.append("#176b5b")
        elif idx == worst_idx:
            bar_colors.append("#c94e3f")
        elif value >= 0:
            bar_colors.append("#9cb9d5")
        else:
            bar_colors.append("#d8bf83")

    fig = go.Figure()
    fig.add_trace(
        go.Bar(
            x=plot["Year"].astype(str),
            y=plot["YoY_Growth_Pct"],
            marker_color=bar_colors,
            text=[f"{value:+.1f}%" for value in plot["YoY_Growth_Pct"]],
            textposition="auto",
            cliponaxis=False,
            name="YoY growth",
            hovertemplate="<b>%{x}</b><br>YoY growth: %{y:+.1f}%<extra></extra>",
        )
    )
    fig.add_trace(
        go.Scatter(
            x=plot["Year"].astype(str),
            y=plot["YoY_Growth_Pct"],
            mode="lines+markers",
            name="Trend",
            line=dict(color="#18222f", width=2),
            marker=dict(size=8, color="#18222f"),
            hovertemplate="<b>%{x}</b><br>Trend: %{y:+.1f}%<extra></extra>",
        )
    )
    fig.add_hline(y=0, line_color="#9aa4ad", line_width=1)
    annotations = []
    for idx, label, color, yshift in [
        (best_idx, "Best year", "#176b5b", 32),
        (worst_idx, "Weakest year", "#c94e3f", -32),
    ]:
        row = plot.loc[idx]
        annotations.append(
            dict(
                x=str(int(row["Year"])),
                y=row["YoY_Growth_Pct"],
                text=label,
                showarrow=True,
                arrowhead=2,
                ax=0,
                ay=-yshift,
                font=dict(color=color, size=12),
                arrowcolor=color,
                bgcolor="#ffffff",
                bordercolor=color,
                borderpad=4,
            )
        )
    fig.update_layout(
        title="Year-to-year growth",
        xaxis_title="Year",
        yaxis_title="Annual growth",
        annotations=annotations,
        showlegend=False,
        bargap=0.28,
    )
    fig.update_yaxes(ticksuffix="%")
    return apply_chart_style(fig, height=420)


def risk_return_scatter(data: pd.DataFrame) -> go.Figure:
    plot = data.dropna(subset=["Total_Risk_Score", "Expected_Growth_2026", "Current_Price_2025"]).copy()
    if plot.empty:
        return go.Figure()
    fig = px.scatter(
        plot,
        x="Total_Risk_Score",
        y="Expected_Growth_2026",
        size="Current_Price_2025",
        color="Investment_Strategy" if "Investment_Strategy" in plot.columns else None,
        hover_name="Suburb",
        hover_data={
            "Current_Price_2025": ":$,.0f",
            "Total_Risk_Score": ":.1f",
            "Expected_Growth_2026": ":.1f",
            "Actual_House_Yield": ":.1f" if "Actual_House_Yield" in plot.columns else False,
        },
        labels=DISPLAY_COLUMN_LABELS,
        color_discrete_sequence=["#176b5b", "#285f96", "#c94e3f", "#a87922", "#5b6573"],
    )
    fig.add_hline(y=plot["Expected_Growth_2026"].median(), line_dash="dot", line_color="#9aa4ad")
    fig.add_vline(x=plot["Total_Risk_Score"].median(), line_dash="dot", line_color="#9aa4ad")
    fig.update_layout(
        title="Risk-return landscape",
        xaxis_title="Total risk score",
        yaxis_title="Expected next-year growth",
    )
    fig.update_yaxes(ticksuffix="%")
    return apply_chart_style(fig, height=520)


def risk_breakdown_chart(row: pd.Series) -> go.Figure:
    scores = pd.DataFrame(
        {
            "Risk area": ["Overall", "Market", "Economic", "Crime"],
            "Score": [
                row.get("Total_Risk_Score"),
                row.get("Market_Risk"),
                row.get("Economic_Risk"),
                row.get("Crime_Rate_Risk"),
            ],
        }
    ).dropna()
    fig = px.bar(
        scores,
        x="Score",
        y="Risk area",
        orientation="h",
        color="Score",
        color_continuous_scale=["#176b5b", "#d9b35f", "#c94e3f"],
        range_color=[0, 100],
    )
    fig.update_layout(
        title="Risk breakdown: shorter bars are safer",
        xaxis_title="Risk score",
        yaxis_title="",
        coloraxis_showscale=False,
    )
    fig.update_xaxes(range=[0, 100])
    return apply_chart_style(fig, height=300)


def rent_comparison_chart(row: pd.Series) -> go.Figure:
    rent = pd.DataFrame(
        {
            "Rent type": ["2021 census rent", "Fair rent estimate", "Market rent estimate"],
            "Weekly rent": [
                row.get("Census_Rent_2021"),
                row.get("Fair_House_Rent_2025"),
                row.get("Actual_House_Rent_2025"),
            ],
        }
    ).dropna()
    fig = px.bar(
        rent,
        x="Rent type",
        y="Weekly rent",
        text=rent["Weekly rent"].apply(fmt_money),
        color="Rent type",
        color_discrete_sequence=["#5b6573", "#176b5b", "#285f96"],
    )
    fig.update_traces(textposition="outside", cliponaxis=False)
    fig.update_layout(
        title="Weekly rent comparison",
        xaxis_title="",
        yaxis_title="Weekly rent",
        showlegend=False,
    )
    fig.update_yaxes(tickprefix="$", tickformat=",.0f")
    return apply_chart_style(fig, height=330)


def crime_composition_chart(row: pd.Series) -> go.Figure:
    data = pd.DataFrame(
        {
            "Crime type": ["Property crimes", "Person crimes"],
            "Count": [
                row.get("Property_Crime_Count"),
                row.get("Person_Crime_Count"),
            ],
        }
    ).dropna()
    fig = px.bar(
        data,
        x="Crime type",
        y="Count",
        text=data["Count"].apply(lambda value: fmt_num(value)),
        color="Crime type",
        color_discrete_sequence=["#285f96", "#c94e3f"],
    )
    fig.update_traces(textposition="outside", cliponaxis=False)
    fig.update_layout(
        title="Crime composition",
        xaxis_title="",
        yaxis_title="Recorded crimes",
        showlegend=False,
    )
    return apply_chart_style(fig, height=320)


def price_distribution(data: pd.DataFrame) -> go.Figure:
    plot = data.dropna(subset=["Current_Price_2025"]).copy()
    fig = px.histogram(
        plot,
        x="Current_Price_2025",
        nbins=26,
        color_discrete_sequence=["#176b5b"],
    )
    fig.update_layout(title="Latest price distribution", xaxis_title="Latest price", yaxis_title="Suburbs")
    fig.update_xaxes(tickprefix="$", tickformat=",.0f")
    return apply_chart_style(fig, height=360)


def ranked_table(data: pd.DataFrame, sort_col: str, columns: list[str], top_n: int = 10, ascending: bool = False) -> pd.DataFrame:
    existing = []
    for col in columns:
        if col in data.columns and col not in existing:
            existing.append(col)
    if sort_col not in data.columns or not existing:
        return pd.DataFrame()
    table = data.dropna(subset=[sort_col]).sort_values(sort_col, ascending=ascending).head(top_n)[existing].copy()
    return table


def format_dashboard_table(table: pd.DataFrame) -> pd.DataFrame:
    out = table.copy()
    for col in out.columns:
        if "Price" in col or "Rent" in col or col in {"Value_Gap"}:
            out[col] = out[col].apply(fmt_price)
        elif "Growth" in col or "Yield" in col or "Percent" in col or "Return" in col:
            out[col] = out[col].apply(fmt_pct)
        elif "Score" in col or "Rate" in col:
            out[col] = out[col].apply(lambda value: fmt_num(value, 1))
    return out.rename(columns=DISPLAY_COLUMN_LABELS)


def color_for_value(value: float, values: pd.Series, palette_name: str) -> str:
    if pd.isna(value):
        return "#aab2b8"
    palette = PALETTES.get(palette_name, PALETTES["price"])
    valid = pd.to_numeric(values, errors="coerce").dropna()
    if valid.empty:
        return palette[2]
    cuts = np.nanpercentile(valid, [20, 40, 60, 80])
    idx = int(np.searchsorted(cuts, value, side="right"))
    return palette[max(0, min(idx, len(palette) - 1))]


def make_map(df: pd.DataFrame, metric_label: str, geojson: dict | None, coords: dict) -> folium.Map:
    metric_col, tooltip_label, palette = METRIC_OPTIONS[metric_label]
    values = pd.to_numeric(df.get(metric_col, pd.Series(dtype=float)), errors="coerce")
    lookup = df.set_index("Suburb").to_dict(orient="index")

    m = folium.Map(location=[-34.9285, 138.6007], zoom_start=11, tiles="CartoDB positron")

    if geojson:
        layer = copy.deepcopy(geojson)
        for feature in layer.get("features", []):
            suburb = feature.get("properties", {}).get("Suburb", "")
            row = lookup.get(suburb, {})
            raw_value = row.get(metric_col, np.nan)
            feature["properties"]["metric_value"] = raw_value
            feature["properties"]["metric_text"] = format_metric_for_col(metric_col, raw_value)
            feature["properties"]["price_text"] = fmt_price(row.get("Current_Price_2025", np.nan))
            feature["properties"]["risk_text"] = safe_text(row.get("Total_Risk_Category", "N/A"))
            feature["properties"]["fill_color"] = color_for_value(raw_value, values, palette)

        folium.GeoJson(
            layer,
            name=metric_label,
            style_function=lambda feature: {
                "fillColor": feature["properties"].get("fill_color", "#aab2b8"),
                "color": "#ffffff",
                "weight": 0.7,
                "fillOpacity": 0.72,
            },
            highlight_function=lambda feature: {
                "weight": 2,
                "color": "#18222f",
                "fillOpacity": 0.9,
            },
            tooltip=folium.GeoJsonTooltip(
                fields=["Suburb", "metric_text", "price_text", "risk_text"],
                aliases=["Suburb", tooltip_label, "Price", "Risk"],
                sticky=True,
                style=(
                    "background-color: white; color: #18222f; font-family: Inter, Arial; "
                    "font-size: 12px; padding: 8px; border: 1px solid #d9dfdc; border-radius: 6px;"
                ),
            ),
        ).add_to(m)
    else:
        for suburb, row in lookup.items():
            if suburb not in coords:
                continue
            raw_value = row.get(metric_col, np.nan)
            color = color_for_value(raw_value, values, palette)
            folium.CircleMarker(
                location=[coords[suburb]["lat"], coords[suburb]["lng"]],
                radius=7,
                color=color,
                fill=True,
                fill_color=color,
                fill_opacity=0.78,
                tooltip=f"{suburb}: {format_metric_for_col(metric_col, raw_value)}",
            ).add_to(m)

    legend = "".join(
        f'<span style="display:inline-block;width:12px;height:12px;background:{color};border-radius:2px;margin-right:6px;"></span>'
        for color in PALETTES.get(palette, PALETTES["price"])
    )
    legend_html = f"""
    <div style="position: fixed; bottom: 24px; left: 24px; z-index: 9999;
        background: white; border: 1px solid #d9dfdc; border-radius: 8px;
        padding: 10px 12px; box-shadow: 0 8px 20px rgba(24,34,47,.15);
        font-family: Inter, Arial; color: #18222f; font-size: 12px;">
        <strong>{escape(metric_label)}</strong><br>
        <span style="color:#647181;">Low</span><br>{legend}<br>
        <span style="color:#647181;">High</span>
    </div>
    """
    m.get_root().html.add_child(folium.Element(legend_html))
    return m


def format_metric_for_col(col: str, value: float | int | None) -> str:
    if col in {"Current_Price_2025", "Forecast_Price_2026"}:
        return fmt_price(value)
    if col in {"Price_Growth_Percent", "Expected_Growth_2026", "Actual_House_Yield"}:
        return fmt_pct(value)
    if col == "Cultural_Diversity_Index":
        return fmt_num(value, 3)
    return fmt_num(value, 1)


def nearest_suburb_from_click(click: dict | None, coords: dict, valid_suburbs: set[str]) -> str | None:
    if not click or not coords:
        return None
    lat = click.get("lat")
    lng = click.get("lng")
    if lat is None or lng is None:
        return None
    nearest = None
    min_dist = float("inf")
    for suburb, coord in coords.items():
        if suburb not in valid_suburbs:
            continue
        dist = (coord["lat"] - lat) ** 2 + (coord["lng"] - lng) ** 2
        if dist < min_dist:
            min_dist = dist
            nearest = suburb
    return nearest


def clicked_suburb_from_map(map_state: dict | None, coords: dict, valid_suburbs: set[str]) -> str | None:
    if not map_state:
        return None

    active = map_state.get("last_active_drawing")
    if isinstance(active, dict):
        props = active.get("properties", {})
        suburb = props.get("Suburb") or props.get("suburb")
        if suburb in valid_suburbs:
            return suburb

    tooltip = map_state.get("last_object_clicked_tooltip")
    if isinstance(tooltip, str):
        suburb = tooltip.split(":")[0].strip()
        if suburb in valid_suburbs:
            return suburb

    return nearest_suburb_from_click(map_state.get("last_clicked"), coords, valid_suburbs)


def pdf_metric(pdf: FPDF, label: str, value: str) -> None:
    pdf.set_font("Helvetica", "B", 9)
    pdf.cell(48, 7, plain_text(label), border=0)
    pdf.set_font("Helvetica", "", 9)
    pdf.cell(0, 7, plain_text(value), new_x=XPos.LMARGIN, new_y=YPos.NEXT)


def pdf_bar(pdf: FPDF, label: str, value: float | int | None, max_value: float, rgb: tuple[int, int, int]) -> None:
    y = pdf.get_y()
    x = pdf.get_x()
    pdf.set_font("Helvetica", "", 8)
    pdf.cell(48, 7, plain_text(label), border=0)
    bar_x = x + 52
    bar_w = 95
    fill_w = 0 if pd.isna(value) else max(0, min(float(value), max_value)) / max_value * bar_w
    pdf.set_draw_color(210, 215, 212)
    pdf.rect(bar_x, y + 1.5, bar_w, 4)
    pdf.set_fill_color(*rgb)
    pdf.rect(bar_x, y + 1.5, fill_w, 4, style="F")
    pdf.set_xy(bar_x + bar_w + 5, y)
    pdf.cell(0, 7, plain_text(fmt_num(value, 1) if not pd.isna(value) else "N/A"), new_x=XPos.LMARGIN, new_y=YPos.NEXT)


def generate_insight_pdf(suburb: str, row: pd.Series, ts: pd.DataFrame) -> bytes:
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Helvetica", "B", 18)
    pdf.set_text_color(24, 34, 47)
    pdf.cell(0, 10, plain_text(f"{suburb} Property Insight Report"), new_x=XPos.LMARGIN, new_y=YPos.NEXT, align="C")
    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(100, 113, 129)
    pdf.cell(0, 7, "Generated from Adelaide Property Decision Dashboard", new_x=XPos.LMARGIN, new_y=YPos.NEXT, align="C")
    pdf.ln(4)

    pdf.set_text_color(24, 34, 47)
    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(0, 8, "Key insights", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf_metric(pdf, "Latest price", fmt_price(row.get("Current_Price_2025")))
    pdf_metric(pdf, "Next-Year Forecast", fmt_price(row.get("Forecast_Price_2026")))
    pdf_metric(pdf, "Next-year growth", fmt_pct(row.get("Expected_Growth_2026")))
    pdf_metric(pdf, "Strategy", str(row.get("Investment_Strategy", "N/A")))
    pdf_metric(pdf, "Risk verdict", risk_plain_language(row))
    pdf_metric(pdf, "Rental verdict", rent_plain_language(row))
    pdf.ln(3)

    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(0, 8, "Visual summary", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_font("Helvetica", "I", 8)
    pdf.cell(0, 6, "Risk bars: shorter is safer. Growth/yield bars: longer is stronger.", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf_bar(pdf, "Risk score", row.get("Total_Risk_Score"), 100, (201, 78, 63))
    pdf_bar(pdf, "Next-year growth", row.get("Expected_Growth_2026"), 100, (23, 107, 91))
    pdf_bar(pdf, "House yield", row.get("Actual_House_Yield"), 8, (168, 121, 34))
    pdf_bar(pdf, "Crime / 1k", row.get("Crime_Rate_Per_1000"), 500, (40, 95, 150))
    pdf.ln(3)

    annual = annual_growth_data(ts, suburb).dropna(subset=["YoY_Growth_Pct"])
    if not annual.empty:
        best = annual.loc[annual["YoY_Growth_Pct"].idxmax()]
        worst = annual.loc[annual["YoY_Growth_Pct"].idxmin()]
        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(0, 8, "Year-to-year growth", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf_metric(pdf, "Best year", f"{int(best['Year'])}: {fmt_pct(best['YoY_Growth_Pct'])}")
        pdf_metric(pdf, "Weakest year", f"{int(worst['Year'])}: {fmt_pct(worst['YoY_Growth_Pct'])}")

    pdf.ln(3)
    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(0, 8, "Crime insights", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf_metric(pdf, "Total crimes", fmt_num(row.get("Total_Crime_Count")))
    pdf_metric(pdf, "Property crimes", fmt_num(row.get("Property_Crime_Count")))
    pdf_metric(pdf, "Person crimes", fmt_num(row.get("Person_Crime_Count")))
    pdf_metric(pdf, "Top offense #1", str(row.get("Crime_Type_1", "N/A")))
    pdf_metric(pdf, "Top offense #2", str(row.get("Crime_Type_2", "N/A")))
    pdf_metric(pdf, "Top offense #3", str(row.get("Crime_Type_3", "N/A")))

    pdf.ln(4)
    pdf.set_font("Helvetica", "I", 8)
    pdf.set_text_color(100, 113, 129)
    pdf.multi_cell(
        0,
        5,
        plain_text(
            "Notes: ABS Census 2021 is used for demographics, income, mortgage, and baseline rent. "
            "Rental values are estimates adjusted from that baseline. Crime statistics cover FY 2019-20 "
            "through Q2 2025-26, including records through 31 December 2025. Model outputs are decision support, not financial advice."
        ),
    )
    return bytes(pdf.output())


def generate_insight_docx(suburb: str, row: pd.Series, ts: pd.DataFrame) -> bytes:
    doc = Document()
    doc.add_heading(f"{suburb} Property Insight Report", 0)
    doc.add_paragraph("Generated from Adelaide Property Decision Dashboard")

    doc.add_heading("Key insights", level=1)
    for label, value in [
        ("Latest price", fmt_price(row.get("Current_Price_2025"))),
        ("Next-Year Forecast", fmt_price(row.get("Forecast_Price_2026"))),
        ("Next-year growth", fmt_pct(row.get("Expected_Growth_2026"))),
        ("Strategy", str(row.get("Investment_Strategy", "N/A"))),
        ("Risk verdict", risk_plain_language(row)),
        ("Rental verdict", rent_plain_language(row)),
    ]:
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(str(value))

    doc.add_heading("Visual summary", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Metric"
    hdr[1].text = "Value"
    hdr[2].text = "Visual bar"
    for label, value, max_value in [
        ("Risk score", row.get("Total_Risk_Score"), 100),
        ("Next-year growth", row.get("Expected_Growth_2026"), 100),
        ("House yield", row.get("Actual_House_Yield"), 8),
        ("Crime / 1k", row.get("Crime_Rate_Per_1000"), 500),
    ]:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = fmt_num(value, 1) if not pd.isna(value) else "N/A"
        cells[2].text = text_bar(value, max_value)

    annual = annual_growth_data(ts, suburb).dropna(subset=["YoY_Growth_Pct"])
    if not annual.empty:
        best = annual.loc[annual["YoY_Growth_Pct"].idxmax()]
        worst = annual.loc[annual["YoY_Growth_Pct"].idxmin()]
        doc.add_heading("Year-to-year growth", level=1)
        doc.add_paragraph(f"Best year: {int(best['Year'])} ({fmt_pct(best['YoY_Growth_Pct'])})")
        doc.add_paragraph(f"Weakest year: {int(worst['Year'])} ({fmt_pct(worst['YoY_Growth_Pct'])})")

    doc.add_heading("Crime insights", level=1)
    for label, value in [
        ("Total crimes", fmt_num(row.get("Total_Crime_Count"))),
        ("Property crimes", fmt_num(row.get("Property_Crime_Count"))),
        ("Person crimes", fmt_num(row.get("Person_Crime_Count"))),
        ("Top offense #1", row.get("Crime_Type_1", "N/A")),
        ("Top offense #2", row.get("Crime_Type_2", "N/A")),
        ("Top offense #3", row.get("Crime_Type_3", "N/A")),
    ]:
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(str(value))

    doc.add_heading("Notes", level=1)
    doc.add_paragraph(
        "ABS Census 2021 is used for demographics, income, mortgage, and baseline rent. "
        "Rental values are estimates adjusted from that baseline. Crime statistics cover FY 2019-20 "
        "through Q2 2025-26, including records through 31 December 2025. Model outputs are decision support, not financial advice."
    )

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def render_rank_cards(data: pd.DataFrame, top_n: int = 5, grid_cols: int = 1) -> None:
    ranked = data.dropna(subset=["Opportunity_Score"]).copy()
    complete = ranked[
        ranked["Forecast_Price_2026"].notna()
        & ranked["Expected_Growth_2026"].notna()
        & ranked["Actual_House_Yield"].notna()
    ]
    if len(complete) >= top_n:
        ranked = complete
    known_strategy = ranked[
        ~ranked.get("Investment_Strategy", pd.Series("", index=ranked.index))
        .fillna("")
        .str.contains("Unknown Growth", case=False, regex=False)
    ]
    if len(known_strategy) >= top_n:
        ranked = known_strategy
    ranked = ranked.sort_values("Opportunity_Score", ascending=False).head(top_n)
    if ranked.empty:
        st.info("No opportunity ranking is available for the current filters.")
        return
    columns = st.columns(grid_cols) if grid_cols > 1 else None
    for idx, (_, row) in enumerate(ranked.iterrows()):
        pills = [
            f'<span class="pill teal">{fmt_price(row.get("Current_Price_2025"))}</span>',
            f'<span class="pill">{fmt_pct(row.get("Expected_Growth_2026"))} next-year growth</span>',
            f'<span class="pill amber">{fmt_pct(row.get("Actual_House_Yield"))} yield</span>',
            f'<span class="pill coral">{safe_text(row.get("Total_Risk_Category", "N/A"))}</span>',
        ]
        html = f"""
        <div class="rank-card">
            <div class="rank-title">
                <strong>{safe_text(row.get("Suburb"))}</strong>
                <span class="rank-score">{fmt_num(row.get("Opportunity_Score"), 1)}</span>
            </div>
            <div class="rank-meta">{safe_text(row.get("Investment_Strategy", "N/A"))}</div>
            <div class="pill-row">{''.join(pills)}</div>
        </div>
        """
        if columns:
            columns[idx % grid_cols].markdown(html, unsafe_allow_html=True)
        else:
            st.markdown(html, unsafe_allow_html=True)


def render_overview(df: pd.DataFrame, ts: pd.DataFrame, view: pd.DataFrame) -> None:
    valid = view[view["Current_Price_2025"].notna()].copy()
    title_block(
        "Adelaide Property Decision Dashboard",
        "Updated 23 April 2026 with 414 suburbs, 29 property quarters, Q1 2026 latest property prices, and crime records through 31 December 2025.",
    )

    c1, c2, c3, c4, c5 = st.columns(5)
    c1.markdown(card("Suburbs loaded", fmt_num(len(df)), "Full master dataset", "teal"), unsafe_allow_html=True)
    c2.markdown(card("Full-price suburbs", fmt_num(valid["Current_Price_2025"].notna().sum()), "Latest price available", "blue"), unsafe_allow_html=True)
    c3.markdown(card("Median price", fmt_price(valid["Current_Price_2025"].median()), "Across active filters", "amber"), unsafe_allow_html=True)
    c4.markdown(card("Median 29-quarter growth", fmt_pct(valid["Price_Growth_Percent"].median()), "Q1 2019 to Q1 2026", "teal"), unsafe_allow_html=True)
    c5.markdown(card("Median risk score", fmt_num(valid["Total_Risk_Score"].median(), 1), "Lower is safer", "coral"), unsafe_allow_html=True)

    section("Quick searches", "Click once to surface common suburb lists.")
    if "overview_action" not in st.session_state:
        st.session_state.overview_action = "Best opportunities"

    actions = [
        ("Best opportunities", "Opportunity_Score", False, "Investment shortlist"),
        ("Most growing", "Price_Growth_Percent", False, "Highest Q1 2019 to Q1 2026 growth"),
        ("Highest next-year growth", "Expected_Growth_2026", False, "Highest next-year model growth"),
        ("Best yield", "Actual_House_Yield", False, "Highest estimated house yield"),
        ("Lowest risk", "Total_Risk_Score", True, "Lowest total risk score"),
        ("Highest crime", "Crime_Rate_Per_1000", False, "Highest crimes per 1,000 people"),
        ("Safest suburbs", "Crime_Rate_Per_1000", True, "Lowest crimes per 1,000 people"),
        ("Most Indian", "Indian_Percent", False, "Highest Indian community share"),
    ]
    action_cols = st.columns(4)
    for idx, (label, _, _, _) in enumerate(actions):
        if action_cols[idx % 4].button(label, use_container_width=True):
            st.session_state.overview_action = label

    selected_action = next((item for item in actions if item[0] == st.session_state.overview_action), actions[0])
    action_label, sort_col, ascending, action_note = selected_action
    section(action_label, action_note)
    action_table = ranked_table(
        valid,
        sort_col,
        [
            "Suburb",
            "Current_Price_2025",
            sort_col,
            "Expected_Growth_2026",
            "Actual_House_Yield",
            "Crime_Rate_Per_1000",
            "Total_Risk_Category",
            "Investment_Strategy",
        ],
        12,
        ascending=ascending,
    )
    st.dataframe(format_dashboard_table(action_table), use_container_width=True, hide_index=True)

    section("Executive snapshot", "The first screen emphasizes investable signals instead of raw tables.")
    render_rank_cards(valid, 6, grid_cols=3)

    section("Risk-return landscape", "Bubble size represents latest price; position shows risk versus expected next-year growth.")
    st.plotly_chart(risk_return_scatter(valid), use_container_width=True)

    left, right = st.columns([1, 1])
    with left:
        st.plotly_chart(price_distribution(valid), use_container_width=True)
    with right:
        top_growth = ranked_table(
            valid,
            "Price_Growth_Percent",
            ["Suburb", "Current_Price_2025", "Price_Growth_Percent", "Expected_Growth_2026", "Total_Risk_Category"],
            12,
        )
        section("Top growth shortlist")
        st.dataframe(format_dashboard_table(top_growth), use_container_width=True, hide_index=True)


def render_explore(df: pd.DataFrame, ts: pd.DataFrame, suburb: str) -> None:
    row_df = df[df["Suburb"] == suburb]
    if row_df.empty:
        st.warning("No data found for the selected suburb.")
        return
    row = row_df.iloc[0]
    title_block(
        f"{suburb.title()} suburb profile",
        "A focused report that turns the output files into a clearer investment story.",
        "Explore suburb",
    )

    strategy = safe_text(row.get("Investment_Strategy", "N/A"))
    risk = safe_text(row.get("Total_Risk_Category", "N/A"))
    st.markdown(
        f"""
        <div class="insight-box">
            <strong>Positioning:</strong> {safe_text(suburb)} has a latest price of
            <strong>{fmt_price(row.get("Current_Price_2025"))}</strong>, with
            <strong>{fmt_pct(row.get("Price_Growth_Percent"))}</strong> growth from Q1 2019 to Q1 2026 and
            a next-year forecast of <strong>{fmt_price(row.get("Forecast_Price_2026"))}</strong>.
            The model classifies it as <strong>{risk}</strong> with strategy:
            <strong>{strategy}</strong>.
        </div>
        """,
        unsafe_allow_html=True,
    )

    c1, c2, c3, c4 = st.columns(4)
    c1.markdown(card("Latest price", fmt_price(row.get("Current_Price_2025")), "Q1 2026 median", "teal"), unsafe_allow_html=True)
    c2.markdown(card("Next-Year Forecast", fmt_price(row.get("Forecast_Price_2026")), fmt_pct(row.get("Expected_Growth_2026")), "blue"), unsafe_allow_html=True)
    c3.markdown(card("House yield", fmt_pct(row.get("Actual_House_Yield")), safe_text(row.get("Affordability_Category", "N/A")), "amber"), unsafe_allow_html=True)
    c4.markdown(card("Risk score", fmt_num(row.get("Total_Risk_Score"), 1), safe_text(row.get("Total_Risk_Category", "N/A")), "coral"), unsafe_allow_html=True)

    tab_growth, tab_risk, tab_crime, tab_people, tab_download = st.tabs([
        "Growth and forecast",
        "Risk and rental",
        "Crime insights",
        "People and community",
        "Downloads",
    ])
    with tab_growth:
        st.plotly_chart(make_price_history(ts, [suburb], df), use_container_width=True)
        annual = annual_growth_data(ts, suburb)
        yoy_fig = yoy_growth_chart(ts, suburb)
        if yoy_fig is not None and not annual.empty:
            annual_plot = annual.dropna(subset=["YoY_Growth_Pct"])
            if not annual_plot.empty:
                best = annual_plot.loc[annual_plot["YoY_Growth_Pct"].idxmax()]
                worst = annual_plot.loc[annual_plot["YoY_Growth_Pct"].idxmin()]
                best_col, worst_col = st.columns(2)
                best_col.markdown(
                    card(
                        "Best growth year",
                        str(int(best["Year"])),
                        f"{fmt_pct(best['YoY_Growth_Pct'])} | {fmt_price_change(best['YoY_Growth_Amount'])}",
                        "teal",
                    ),
                    unsafe_allow_html=True,
                )
                worst_col.markdown(
                    card(
                        "Weakest growth year",
                        str(int(worst["Year"])),
                        f"{fmt_pct(worst['YoY_Growth_Pct'])} | {fmt_price_change(worst['YoY_Growth_Amount'])}",
                        "coral",
                    ),
                    unsafe_allow_html=True,
                )
            st.plotly_chart(yoy_fig, use_container_width=True)
        subcols = st.columns(3)
        subcols[0].markdown(card("Price volatility", fmt_num(row.get("Price_Volatility"), 0), "Std deviation", "blue"), unsafe_allow_html=True)
        gap_value, gap_note = value_gap_note(row.get("Value_Gap_Percent"), row.get("Value_Category", "N/A"))
        subcols[1].markdown(card("Model price gap", gap_value, gap_note, "teal"), unsafe_allow_html=True)
        subcols[2].markdown(card("Prediction error", fmt_pct(row.get("Error_Percent")), "Latest model check", "amber"), unsafe_allow_html=True)
        st.markdown(
            """
            <div class="note-box">
                Model price gap compares the latest price with the model's predicted/fair price.
                Negative means the actual price is above the model estimate; positive means it is below the model estimate.
                Small negative or positive gaps can still be classified as fairly valued.
            </div>
            """,
            unsafe_allow_html=True,
        )

    with tab_risk:
        st.markdown(
            f"""
            <div class="insight-box">
                <strong>Risk verdict:</strong> {safe_text(risk_plain_language(row))}
            </div>
            """,
            unsafe_allow_html=True,
        )

        c1, c2, c3, c4 = st.columns(4)
        c1.markdown(card("Overall risk", safe_text(row.get("Total_Risk_Category", "N/A")), f"Score {fmt_num(row.get('Total_Risk_Score'), 1)}", "coral"), unsafe_allow_html=True)
        c2.markdown(card("Market risk", safe_text(row.get("Market_Risk_Category", "N/A")), "Price + growth stability", "amber"), unsafe_allow_html=True)
        c3.markdown(card("Crime risk", safe_text(row.get("Crime_Risk_Category", "N/A")), f"{fmt_num(row.get('Crime_Rate_Per_1000'), 1)} crimes / 1k", "blue"), unsafe_allow_html=True)
        c4.markdown(card("Economic risk", safe_text(row.get("Economic_Risk_Category", "N/A")), "Income + employment signal", "teal"), unsafe_allow_html=True)

        st.markdown(
            f"""
            <div class="insight-box">
                <strong>Rental verdict:</strong> {safe_text(rent_plain_language(row))}
            </div>
            """,
            unsafe_allow_html=True,
        )

        c1, c2, c3, c4 = st.columns(4)
        c1.markdown(card("Fair rent estimate", fmt_money(row.get("Fair_House_Rent_2025")), "Weekly house rent", "teal"), unsafe_allow_html=True)
        c2.markdown(card("Market rent estimate", fmt_money(row.get("Actual_House_Rent_2025")), "Estimated weekly rent", "blue"), unsafe_allow_html=True)
        c3.markdown(card("Rent above fair", fmt_pct(row.get("Greediness_Percent")), "Actual estimate vs fair estimate", "coral"), unsafe_allow_html=True)
        c4.markdown(card("House yield", fmt_pct(row.get("Actual_House_Yield")), safe_text(row.get("Affordability_Category", "N/A")), "amber"), unsafe_allow_html=True)

        viz_left, viz_right = st.columns([1, 1])
        with viz_left:
            st.plotly_chart(risk_breakdown_chart(row), use_container_width=True)
        with viz_right:
            st.plotly_chart(rent_comparison_chart(row), use_container_width=True)

        st.markdown(
            """
            <div class="note-box">
                Plain-English reading: lower risk categories are safer; higher yield is better for rental return.
                The rent-above-fair figure is not a moral judgement. It shows how much the estimated market rent
                is above the inflation-adjusted census rent benchmark.
            </div>
            """,
            unsafe_allow_html=True,
        )

    with tab_crime:
        st.markdown(
            f"""
            <div class="insight-box">
                <strong>Crime summary:</strong> {safe_text(suburb)} has
                <strong>{fmt_num(row.get("Total_Crime_Count"))}</strong> total crimes in the main suburb dataset,
                with a rate of <strong>{fmt_num(row.get("Crime_Rate_Per_1000"), 1)}</strong> crimes per 1,000 people.
            </div>
            """,
            unsafe_allow_html=True,
        )

        c1, c2, c3, c4 = st.columns(4)
        crime_rate = row.get("Crime_Rate_Per_1000")
        crime_label = "Low" if not pd.isna(crime_rate) and crime_rate < 50 else "Moderate" if not pd.isna(crime_rate) and crime_rate < 150 else "High"
        c1.markdown(card("Total crimes", fmt_num(row.get("Total_Crime_Count")), "Main suburb dataset", "coral"), unsafe_allow_html=True)
        c2.markdown(card("Property crimes", fmt_num(row.get("Property_Crime_Count")), "Theft, damage, related", "blue"), unsafe_allow_html=True)
        c3.markdown(card("Person crimes", fmt_num(row.get("Person_Crime_Count")), "Against-the-person offenses", "amber"), unsafe_allow_html=True)
        c4.markdown(card("Crime / 1,000 ppl", fmt_num(crime_rate, 1), crime_label, "teal" if crime_label == "Low" else "amber" if crime_label == "Moderate" else "coral"), unsafe_allow_html=True)

        left, right = st.columns([1, 1])
        with left:
            st.plotly_chart(crime_composition_chart(row), use_container_width=True)
        with right:
            no_of_crimes = row.get("No_of_Crimes")
            st.markdown(
                f"""
                <div class="rank-card">
                    <div class="rank-title">
                        <strong>Top offense types</strong>
                        <span class="rank-score">{fmt_num(no_of_crimes)}</span>
                    </div>
                    <div class="rank-meta">Total recorded in offense analysis</div>
                    <div class="pill-row">
                        <span class="pill coral">#1 {safe_text(row.get("Crime_Type_1", "N/A"))}</span>
                        <span class="pill amber">#2 {safe_text(row.get("Crime_Type_2", "N/A"))}</span>
                        <span class="pill">#3 {safe_text(row.get("Crime_Type_3", "N/A"))}</span>
                    </div>
                </div>
                """,
                unsafe_allow_html=True,
            )

        st.markdown(
            """
            <div class="note-box">
                Crime Data Note: Crime statistics cover SA Government records from FY 2019-20 through
                Q2 2025-26, including records through 31 December 2025. Counts are cumulative across reporting
                periods, and crime rate per 1,000 people helps compare suburbs with different population sizes.
            </div>
            """,
            unsafe_allow_html=True,
        )

    with tab_people:
        left, right = st.columns([1, 1])
        with left:
            people = pd.DataFrame(
                {
                    "Metric": ["Population", "Median age", "HH income / week", "Mortgage / month"],
                    "Value": [
                        row.get("G01_Population_Total"),
                        row.get("G02_Median_age_persons"),
                        row.get("G02_Median_tot_hhd_inc_weekly"),
                        row.get("G02_Median_mortgage_repay_monthly"),
                    ],
                }
            )
            st.dataframe(format_dashboard_table(people), use_container_width=True, hide_index=True)
        with right:
            communities = ["Indian", "Chinese", "Vietnamese", "Italian", "Greek"]
            culture = pd.DataFrame(
                {
                    "Community": communities,
                    "Percent": [row.get(f"{name}_Percent") for name in communities],
                    "People": [row.get(f"{name}_Population") for name in communities],
                }
            ).dropna()
            culture["Label"] = culture.apply(
                lambda item: f"{item['Percent']:.1f}%<br>{fmt_num(item['People'])} people",
                axis=1,
            )
            fig = px.bar(culture, x="Community", y="Percent", color="Community", color_discrete_sequence=["#176b5b", "#285f96", "#c94e3f", "#a87922", "#5b6573"])
            fig.update_traces(text=culture["Label"], textposition="outside", cliponaxis=False)
            fig.update_layout(title="Community profile", showlegend=False, yaxis_title="Share of population")
            fig.update_yaxes(ticksuffix="%")
            st.plotly_chart(apply_chart_style(fig, 340), use_container_width=True)

        st.markdown(
            """
            <div class="note-box">
                Census and rent note: population, age, household income, mortgage, cultural demographics,
                and baseline rent values come from ABS Census 2021. Rental figures shown elsewhere are
                estimates adjusted forward from that 2021 census baseline.
            </div>
            """,
            unsafe_allow_html=True,
        )

    with tab_download:
        section("Download insights", "Export this suburb's key findings, plain-English notes, and simple visual summaries.")
        st.markdown(
            """
            <div class="note-box">
                The PDF and Word reports include the main suburb metrics, risk/rental explanation,
                year-to-year best and weakest growth years, crime insights, and simple visual bars.
            </div>
            """,
            unsafe_allow_html=True,
        )
        d1, d2, _ = st.columns([1, 1, 2])
        with d1:
            st.download_button(
                "Download PDF report",
                data=generate_insight_pdf(suburb, row, ts),
                file_name=f"{suburb}_property_insights.pdf",
                mime="application/pdf",
                use_container_width=True,
            )
        with d2:
            st.download_button(
                "Download Word report",
                data=generate_insight_docx(suburb, row, ts),
                file_name=f"{suburb}_property_insights.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )


def render_compare(df: pd.DataFrame, ts: pd.DataFrame, suburbs: list[str]) -> None:
    title_block(
        "Compare suburbs",
        "Compare price paths, next-year forecast, risk, rent, safety, and demographics side by side.",
        "Comparison workflow",
    )
    if not suburbs:
        st.info("Choose at least one suburb from the sidebar comparison control.")
        return
    selected = df[df["Suburb"].isin(suburbs)].copy()
    st.plotly_chart(make_price_history(ts, suburbs, None), use_container_width=True)

    metrics = [
        "Suburb",
        "Current_Price_2025",
        "Price_Growth_Percent",
        "Expected_Growth_2026",
        "Total_Risk_Score",
        "Actual_House_Yield",
        "Crime_Rate_Per_1000",
        "G02_Median_tot_hhd_inc_weekly",
        "Investment_Strategy",
    ]
    section("Comparison table")
    st.dataframe(format_dashboard_table(selected[[col for col in metrics if col in selected.columns]]), use_container_width=True, hide_index=True)

    long_metrics = selected[
        [
            col
            for col in ["Suburb", "Price_Growth_Percent", "Expected_Growth_2026", "Actual_House_Yield", "Total_Risk_Score"]
            if col in selected.columns
        ]
    ].melt(id_vars="Suburb", var_name="Metric", value_name="Value")
    long_metrics["Metric"] = long_metrics["Metric"].map(DISPLAY_COLUMN_LABELS).fillna(long_metrics["Metric"])
    fig = px.bar(long_metrics, x="Suburb", y="Value", color="Metric", barmode="group", color_discrete_sequence=["#176b5b", "#285f96", "#c94e3f", "#a87922"])
    fig.update_layout(title="Signal comparison", xaxis_title="", yaxis_title="Score / percent")
    st.plotly_chart(apply_chart_style(fig, 430), use_container_width=True)


def render_opportunities(df: pd.DataFrame) -> None:
    title_block(
        "Opportunity finder",
        "Rank suburbs using model score, risk, next-year growth, rental yield, and affordability.",
        "Decision workflow",
    )
    valid = df[df["Current_Price_2025"].notna()].copy()
    st.markdown(
        """
        <div class="note-box">
            <strong>How to use this:</strong> set the maximum price you can afford, then choose the minimum
            future growth and rental yield you want. Next-year growth is the model's expected price increase.
            House yield is estimated yearly rent divided by the latest property price, so higher usually means
            better rental return.
        </div>
        """,
        unsafe_allow_html=True,
    )
    with st.container(border=True):
        controls = st.columns([1, 1, 1, 0.8])
        max_budget = controls[0].number_input("Maximum budget", min_value=0, value=1_200_000, step=50_000)
        min_growth = controls[1].number_input(
            "Minimum next-year growth %",
            value=0.0,
            step=1.0,
            help="Only show suburbs where the model expects at least this much next-year price growth.",
        )
        min_yield = controls[2].number_input(
            "Minimum house yield %",
            value=0.0,
            step=0.25,
            help="Only show suburbs where estimated yearly rent is at least this percentage of the property price.",
        )
        hide_high_risk = controls[3].toggle("Hide high risk", value=True)
        st.caption(
            "Example: 5% next-year growth means the model expects the price to rise by at least 5%. "
            "A 3.5% house yield means annual rent is about 3.5% of the property price."
        )

    filtered = valid[
        (valid["Current_Price_2025"] <= max_budget)
        & (valid["Expected_Growth_2026"].fillna(-999) >= min_growth)
        & (valid["Actual_House_Yield"].fillna(-999) >= min_yield)
    ].copy()
    if hide_high_risk and "Total_Risk_Category" in filtered.columns:
        filtered = filtered[~filtered["Total_Risk_Category"].isin(["High Risk", "Very High Risk"])]

    c1, c2, c3, c4 = st.columns(4)
    c1.markdown(card("Matches", fmt_num(len(filtered)), "After filters", "teal"), unsafe_allow_html=True)
    c2.markdown(card("Median price", fmt_price(filtered["Current_Price_2025"].median()), "Filtered suburbs", "blue"), unsafe_allow_html=True)
    c3.markdown(card("Median next-year growth", fmt_pct(filtered["Expected_Growth_2026"].median()), "Model expectation", "amber"), unsafe_allow_html=True)
    c4.markdown(card("Median yield", fmt_pct(filtered["Actual_House_Yield"].median()), "House yield", "coral"), unsafe_allow_html=True)

    section("Best matches")
    render_rank_cards(filtered, 6, grid_cols=3)

    section("Risk-return view")
    st.plotly_chart(risk_return_scatter(filtered), use_container_width=True)

    section("Ranked shortlist")
    table = ranked_table(
        filtered,
        "Opportunity_Score",
        [
            "Suburb",
            "Current_Price_2025",
            "Opportunity_Score",
            "Expected_Growth_2026",
            "Actual_House_Yield",
            "Total_Risk_Score",
            "Total_Risk_Category",
            "Investment_Strategy",
        ],
        30,
    )
    st.dataframe(format_dashboard_table(table), use_container_width=True, hide_index=True)


def render_map_lab(df: pd.DataFrame, ts: pd.DataFrame, geojson: dict | None, coords: dict) -> None:
    title_block(
        "Metric map lab",
        "Switch the suburb map between price, growth, risk, yield, crime, and cultural diversity.",
        "Spatial workflow",
    )
    st.markdown(
        """
        <div class="map-control-title">Choose what the map colours mean</div>
        <div class="map-control-help">Choose map metric below to switch between price, growth, risk, yield, crime, and diversity.</div>
        """,
        unsafe_allow_html=True,
    )
    metric_label = st.selectbox("Choose map metric", list(METRIC_OPTIONS.keys()), index=0, label_visibility="collapsed")
    st.caption("Click any suburb on the map to open its full detail panel below.")
    map_state = st_folium(
        make_map(df, metric_label, geojson, coords),
        height=620,
        use_container_width=True,
        key=f"metric_map_{metric_label}",
    )

    valid_suburbs = set(df["Suburb"].dropna().unique().tolist())
    clicked = clicked_suburb_from_map(map_state, coords, valid_suburbs)
    if clicked:
        st.session_state.map_detail_suburb = clicked

    col, detail_col = st.columns([1.1, 1.2])
    metric_col = METRIC_OPTIONS[metric_label][0]
    with col:
        section("Highest suburbs for selected metric")
        table = ranked_table(df, metric_col, ["Suburb", metric_col, "Current_Price_2025", "Total_Risk_Category"], 12)
        st.dataframe(format_dashboard_table(table), use_container_width=True, hide_index=True)
    with detail_col:
        selected = st.session_state.get("map_detail_suburb")
        if selected in valid_suburbs:
            row = df[df["Suburb"] == selected].iloc[0]
            section("Selected suburb")
            st.markdown(
                f"""
                <div class="insight-box">
                    <strong>{safe_text(selected)}</strong><br>
                    Latest price: <strong>{fmt_price(row.get("Current_Price_2025"))}</strong> |
                    Growth: <strong>{fmt_pct(row.get("Price_Growth_Percent"))}</strong> |
                    Risk: <strong>{safe_text(row.get("Total_Risk_Category", "N/A"))}</strong><br>
                    Strategy: <strong>{safe_text(row.get("Investment_Strategy", "N/A"))}</strong>
                </div>
                """,
                unsafe_allow_html=True,
            )
            c1, c2, c3 = st.columns(3)
            c1.markdown(card("Next-Year Forecast", fmt_price(row.get("Forecast_Price_2026")), fmt_pct(row.get("Expected_Growth_2026")), "blue"), unsafe_allow_html=True)
            c2.markdown(card("Yield", fmt_pct(row.get("Actual_House_Yield")), safe_text(row.get("Affordability_Category", "N/A")), "amber"), unsafe_allow_html=True)
            c3.markdown(card("Crime / 1k", fmt_num(row.get("Crime_Rate_Per_1000"), 1), safe_text(row.get("Crime_Risk_Category", "N/A")), "coral"), unsafe_allow_html=True)
        else:
            section("Selected suburb")
            st.info("Click a suburb on the map to show its details here.")

    selected = st.session_state.get("map_detail_suburb")
    if selected in valid_suburbs:
        st.markdown("---")
        render_explore(df, ts, selected)


def render_methodology(df: pd.DataFrame, ts: pd.DataFrame) -> None:
    title_block(
        "Methodology and coverage",
        "What this app uses, where the outputs come from, and how to explain missing values.",
        "Project notes",
    )
    c1, c2, c3, c4 = st.columns(4)
    c1.markdown(card("Master suburbs", fmt_num(len(df)), "Base analysis rows", "teal"), unsafe_allow_html=True)
    c2.markdown(card("Time series rows", fmt_num(len(ts)), "29 property quarters", "blue"), unsafe_allow_html=True)
    c3.markdown(card("Prediction rows", fmt_num(df["Predicted_Price_2025"].notna().sum()), "Suburbs with ML output", "amber"), unsafe_allow_html=True)
    c4.markdown(card("Rental rows", fmt_num(df["Actual_House_Yield"].notna().sum()), "Suburbs with rental analysis", "coral"), unsafe_allow_html=True)

    st.markdown(
        """
        <div class="note-box">
            Latest refresh: 23 April 2026. Property prices cover Q1 2019 to Q1 2026 across 29 quarters.
            Demographics and rental inputs are based on ABS Census 2021, and rental figures are inflation-adjusted
            estimates. Crime statistics cover SA Government records from FY 2019-20 through Q2 2025-26, including
            records through 31 December 2025. The next-year forecast and risk outputs are model outputs from the
            analysis project and should be treated as decision support, not financial advice.
        </div>
        """,
        unsafe_allow_html=True,
    )

    section("Data coverage")
    coverage = pd.DataFrame(
        {
            "Output area": [
                "Latest price",
                "Prediction and risk",
                "Rental and yield",
                "Cultural demographics",
                "Crime offense analysis",
            ],
            "Available suburbs": [
                df["Current_Price_2025"].notna().sum(),
                df["Predicted_Price_2025"].notna().sum(),
                df["Actual_House_Yield"].notna().sum(),
                df["Cultural_Diversity_Index"].notna().sum(),
                df["Crime_Type_1"].notna().sum(),
            ],
            "Total suburbs": [len(df)] * 5,
        }
    )
    coverage["Coverage %"] = coverage["Available suburbs"] / coverage["Total suburbs"] * 100
    st.dataframe(format_dashboard_table(coverage), use_container_width=True, hide_index=True)


def render_glossary() -> None:
    title_block(
        "Glossary",
        "Plain-English meanings for the property, risk, rental, and model terms used in this dashboard.",
        "Help",
    )
    terms = [
        ("Latest price", "The suburb's latest median property price from the processed property sales data, with Q1 2026 as the latest property quarter."),
        ("Next-year growth", "The model's expected percentage price increase over the next-year forecast period. It is a prediction, not a guarantee."),
        ("House yield", "Estimated yearly rent divided by property price. Higher yield usually means better rental income return."),
        ("Risk score", "A combined model score using market, crime, economic, prediction, and growth signals. Lower is safer."),
        ("Market risk", "Risk from price level, growth pattern, volatility, and market movement."),
        ("Crime risk", "Risk based on recorded crime rates. It is not a full personal safety rating."),
        ("Economic risk", "Risk based on local income, employment, and education signals."),
        ("Fair rent estimate", "An inflation-adjusted rent estimate using ABS Census 2021 rent as the baseline."),
        ("Market rent estimate", "Estimated current rent after adjusting the baseline rent upward."),
        ("Rent above fair", "How far the market rent estimate is above the fair-rent estimate. It is a model comparison, not a moral judgement."),
        ("Affordability category", "How stressful rent looks compared with income. Severe crisis means rent is high relative to income."),
        ("Model price gap", "Latest price compared with the model's predicted/fair price. Negative means actual is above the model estimate."),
        ("Fairly valued", "The actual price is close enough to the model estimate that it is not strongly overvalued or undervalued."),
        ("Opportunity score", "A ranking score combining growth, value, safety, yield, and risk. It helps shortlist suburbs."),
        ("Cultural diversity index", "A summary measure of cultural mix from Census 2021 community data. Higher means more diverse."),
        ("Crime / 1k", "Recorded crimes per 1,000 residents. This helps compare suburbs with different population sizes."),
        ("Prediction error", "How far the model's predicted latest price was from the latest observed price."),
        ("Price volatility", "How much prices have moved around over time. Higher volatility means less stable price history."),
    ]
    left, right = st.columns([1, 1])
    for idx, (term, meaning) in enumerate(terms):
        html = f"""
        <div class="glossary-term">
            <strong>{escape(term)}</strong>
            <span>{escape(meaning)}</span>
        </div>
        """
        if idx % 2 == 0:
            left.markdown(html, unsafe_allow_html=True)
        else:
            right.markdown(html, unsafe_allow_html=True)


def main() -> None:
    df, ts = load_data()
    geojson = load_geojson()
    coords = load_coordinates()

    if df.empty:
        st.error("No data loaded. Check the data folder paths.")
        return

    suburbs = sorted(df["Suburb"].dropna().unique().tolist())
    with st.sidebar:
        st.markdown("### Adelaide dashboard")
        st.caption("Presentation-focused Streamlit app")
        page = st.radio(
            "View",
            ["Overview", "Explore suburb", "Compare suburbs", "Opportunity finder", "Map lab", "Methodology", "Glossary"],
            label_visibility="collapsed",
        )
        st.divider()
        selected_suburb = st.selectbox("Suburb", suburbs, index=suburbs.index("ADELAIDE") if "ADELAIDE" in suburbs else 0)
        compare_suburbs = st.multiselect(
            "Compare suburbs",
            suburbs,
            default=[name for name in ["ADELAIDE", "GLENELG", "SALISBURY", "NORWOOD"] if name in suburbs],
            max_selections=6,
        )

        st.divider()
        st.caption("This file is separate from app.py and only reads existing CSV/GeoJSON outputs.")

    if page == "Overview":
        render_overview(df, ts, df)
    elif page == "Explore suburb":
        render_explore(df, ts, selected_suburb)
    elif page == "Compare suburbs":
        render_compare(df, ts, compare_suburbs)
    elif page == "Opportunity finder":
        render_opportunities(df)
    elif page == "Map lab":
        render_map_lab(df, ts, geojson, coords)
    elif page == "Methodology":
        render_methodology(df, ts)
    else:
        render_glossary()


if __name__ == "__main__":
    main()
