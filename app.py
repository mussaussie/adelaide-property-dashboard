"""
Adelaide Property Market Explorer — Streamlit Web App
======================================================
Interactive dashboard for 414 Adelaide suburbs with demographics,
crime, growth, predictions, rental, risk & cultural community data.

Author: Abdul Mussavir | February 2026
"""

import streamlit as st
import pandas as pd
import numpy as np
import json
import os
import io
import plotly.graph_objects as go
import plotly.express as px
import folium
from streamlit_folium import st_folium
from fpdf import FPDF
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ─────────────────────────────────────────────────────────────────────────────
# PAGE CONFIG
# ─────────────────────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="Adelaide Property Market Explorer",
    page_icon="🏡",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ─────────────────────────────────────────────────────────────────────────────
# CUSTOM CSS — Vibrant modern dark theme
# ─────────────────────────────────────────────────────────────────────────────
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800&display=swap');

    html, body, [class*="css"] { font-family: 'Inter', sans-serif; font-size: 16px; }

    /* Animated gradient header */
    .main-header {
        background: linear-gradient(135deg, #0f0c29, #302b63, #24243e, #1a0a3e);
        background-size: 300% 300%;
        animation: headerShift 8s ease infinite;
        padding: 2.5rem 2.5rem;
        border-radius: 18px;
        margin-bottom: 1.5rem;
        color: white;
        text-align: center;
        box-shadow: 0 8px 32px rgba(80, 60, 200, 0.25);
        position: relative;
        overflow: hidden;
    }
    @keyframes headerShift {
        0% { background-position: 0% 50%; }
        50% { background-position: 100% 50%; }
        100% { background-position: 0% 50%; }
    }
    .main-header::before {
        content: '';
        position: absolute;
        top: -50%;
        left: -50%;
        width: 200%;
        height: 200%;
        background: radial-gradient(circle, rgba(120,100,255,0.08) 0%, transparent 70%);
        animation: pulse 4s ease-in-out infinite;
    }
    @keyframes pulse {
        0%, 100% { transform: scale(1); opacity: 0.5; }
        50% { transform: scale(1.1); opacity: 1; }
    }
    .main-header h1 {
        font-size: 2.4rem;
        font-weight: 800;
        margin: 0;
        position: relative;
        background: linear-gradient(90deg, #f093fb, #f5576c, #ffd86f, #f093fb);
        background-size: 200% auto;
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        animation: shimmer 3s linear infinite;
    }
    @keyframes shimmer {
        to { background-position: 200% center; }
    }
    .main-header p {
        color: #c8c8e8;
        font-size: 1.05rem;
        margin-top: 0.5rem;
        position: relative;
    }

    /* Metric cards with gradient borders */
    .metric-card {
        background: linear-gradient(145deg, #1a1a2e, #16213e);
        border: 1px solid transparent;
        border-image: linear-gradient(135deg, #667eea33, #764ba233) 1;
        border-radius: 14px;
        padding: 1.2rem 1.4rem;
        text-align: center;
        transition: all 0.35s cubic-bezier(0.4, 0, 0.2, 1);
        box-shadow: 0 4px 15px rgba(0,0,0,0.2);
        position: relative;
        overflow: hidden;
    }
    .metric-card::before {
        content: '';
        position: absolute;
        top: 0; left: 0; right: 0;
        height: 3px;
        background: linear-gradient(90deg, #667eea, #764ba2, #f093fb);
        opacity: 0;
        transition: opacity 0.3s ease;
    }
    .metric-card:hover {
        transform: translateY(-4px);
        box-shadow: 0 12px 35px rgba(100,100,255,0.2);
    }
    .metric-card:hover::before {
        opacity: 1;
    }
    .metric-card .label {
        color: #9999bb;
        font-size: 0.85rem;
        font-weight: 600;
        text-transform: uppercase;
        letter-spacing: 0.8px;
    }
    .metric-card .value {
        color: #ffffff;
        font-size: 1.65rem;
        font-weight: 700;
        margin: 0.3rem 0;
    }
    .metric-card .delta {
        font-size: 0.9rem;
        font-weight: 500;
    }
    .delta-pos { color: #69f0ae; }
    .delta-neg { color: #ff5252; }
    .delta-neu { color: #ffd740; }

    /* Section headers with accent line */
    .section-header {
        font-size: 1.45rem;
        font-weight: 700;
        color: #e0e0ff;
        margin: 1.8rem 0 0.8rem 0;
        padding-bottom: 0.5rem;
        border-bottom: 3px solid;
        border-image: linear-gradient(90deg, #667eea, #764ba2, transparent) 1;
    }

    /* Badges */
    .badge {
        display: inline-block;
        padding: 0.3rem 0.8rem;
        border-radius: 20px;
        font-size: 0.85rem;
        font-weight: 600;
        letter-spacing: 0.3px;
    }
    .badge-green { background: rgba(76, 175, 80, 0.2); color: #69f0ae; border: 1px solid #69f0ae44; }
    .badge-orange { background: rgba(255, 152, 0, 0.2); color: #ffd740; border: 1px solid #ffd74044; }
    .badge-red { background: rgba(244, 67, 54, 0.2); color: #ff5252; border: 1px solid #ff525244; }
    .badge-blue { background: rgba(33, 150, 243, 0.2); color: #82b1ff; border: 1px solid #82b1ff44; }

    /* Disclaimer box */
    .disclaimer-box {
        background: linear-gradient(135deg, #1a1a1a, #2d1f1f);
        border-left: 4px solid #ff6f00;
        border-radius: 0 10px 10px 0;
        padding: 1rem 1.2rem;
        margin: 1rem 0;
        font-size: 0.95rem;
        color: #ffcc80;
    }
    .disclaimer-box strong { color: #ffa726; }

    /* Styled download buttons */
    .stDownloadButton > button {
        background: linear-gradient(135deg, #667eea, #764ba2) !important;
        color: white !important;
        border: none !important;
        border-radius: 10px !important;
        padding: 0.5rem 1.5rem !important;
        font-weight: 600 !important;
        transition: all 0.3s ease !important;
        box-shadow: 0 4px 15px rgba(102, 126, 234, 0.3) !important;
    }
    .stDownloadButton > button:hover {
        transform: translateY(-2px) !important;
        box-shadow: 0 6px 20px rgba(102, 126, 234, 0.5) !important;
    }

    /* Sidebar styling */
    div[data-testid="stSidebar"] {
        background: linear-gradient(180deg, #0f0c29, #1a1a2e, #16213e);
    }
    div[data-testid="stSidebar"] .stMarkdown p,
    div[data-testid="stSidebar"] .stMarkdown li,
    div[data-testid="stSidebar"] label {
        color: #c8c8e0;
    }

    /* Tabs styling */
    .stTabs [data-baseweb="tab-list"] {
        gap: 8px;
        background: linear-gradient(90deg, #0f0c29, #1a1a2e);
        padding: 0.5rem;
        border-radius: 12px;
    }
    .stTabs [data-baseweb="tab"] {
        border-radius: 8px;
        color: #9999bb;
        font-weight: 600;
        padding: 0.5rem 1rem;
    }
    .stTabs [aria-selected="true"] {
        background: linear-gradient(135deg, #667eea, #764ba2) !important;
        color: white !important;
    }

    /* Growth table styling */
    .growth-positive { color: #69f0ae; font-weight: 600; }
    .growth-negative { color: #ff5252; font-weight: 600; }

    /* YoY growth card */
    .yoy-card {
        background: linear-gradient(145deg, #1a1a2e, #1e2a4a);
        border: 1px solid #2a3a6a;
        border-radius: 12px;
        padding: 1rem;
        margin: 0.5rem 0;
    }
    .yoy-card .year { color: #82b1ff; font-weight: 700; font-size: 1.2rem; }
    .yoy-card .price { color: #e0e0ff; font-size: 1.4rem; font-weight: 600; }
    .yoy-card .growth-up { color: #69f0ae; font-weight: 600; font-size: 1rem; }
    .yoy-card .growth-down { color: #ff5252; font-weight: 600; font-size: 1rem; }

    /* ── MOBILE RESPONSIVE ── */
    @media screen and (max-width: 768px) {
        html, body, [class*="css"] { font-size: 14px; }

        .main-header { padding: 1.5rem 1rem; border-radius: 12px; }
        .main-header h1 { font-size: 1.5rem; }
        .main-header p { font-size: 0.85rem; }

        .metric-card {
            padding: 0.8rem 0.6rem;
            border-radius: 10px;
        }
        .metric-card .label { font-size: 0.65rem; letter-spacing: 0.3px; }
        .metric-card .value { font-size: 1.1rem; }
        .metric-card .delta { font-size: 0.75rem; }
        .metric-card:hover { transform: none; }  /* disable hover lift on touch */

        .section-header { font-size: 1.1rem; margin: 1rem 0 0.5rem 0; }

        .badge { font-size: 0.7rem; padding: 0.2rem 0.5rem; }

        .disclaimer-box { font-size: 0.8rem; padding: 0.7rem 0.8rem; }

        .yoy-card { padding: 0.7rem; }
        .yoy-card .year { font-size: 1rem; }
        .yoy-card .price { font-size: 1.1rem; }
        .yoy-card .growth-up, .yoy-card .growth-down { font-size: 0.85rem; }

        /* Stack tabs vertically on mobile */
        .stTabs [data-baseweb="tab-list"] {
            flex-wrap: wrap;
            gap: 4px;
            padding: 0.3rem;
        }
        .stTabs [data-baseweb="tab"] {
            font-size: 0.75rem;
            padding: 0.35rem 0.6rem;
            min-width: auto;
        }

        /* Sidebar narrower on mobile */
        div[data-testid="stSidebar"] > div { padding: 0.5rem; }

        /* Make Streamlit columns stack on mobile */
        [data-testid="column"] {
            width: 100% !important;
            flex: 1 1 100% !important;
            min-width: 100% !important;
        }
    }

    /* Tablet breakpoint */
    @media screen and (min-width: 769px) and (max-width: 1024px) {
        .main-header h1 { font-size: 1.9rem; }
        .metric-card .value { font-size: 1.35rem; }
        .metric-card .label { font-size: 0.75rem; }
        .section-header { font-size: 1.25rem; }

        [data-testid="column"] {
            min-width: 45% !important;
        }
    }

    /* Touch device optimizations */
    @media (hover: none) and (pointer: coarse) {
        .metric-card:hover { transform: none; box-shadow: 0 4px 15px rgba(0,0,0,0.2); }
        .metric-card:hover::before { opacity: 0; }
        .stDownloadButton > button:hover { transform: none !important; }
    }
</style>
""", unsafe_allow_html=True)

# ─────────────────────────────────────────────────────────────────────────────
# DATA LOADING — Optimized with usecols where possible
# ─────────────────────────────────────────────────────────────────────────────
BASE = os.path.dirname(os.path.abspath(__file__))


# Columns we actually need from each dataset
_MASTER_CORE_COLS = [
    'Suburb', 'Avg_Price_All_Time', 'Median_Price_All_Time', 'Min_Price_Ever',
    'Max_Price_Ever', 'Price_Volatility', 'Quarter_Count', 'Current_Price_2025',
    'First_Price_2019', 'Price_Growth_Amount', 'Price_Growth_Percent',
    'Total_Crime_Count', 'Crime_OFFENCES_AGAINST_PROPERTY',
    'Crime_OFFENCES_AGAINST_THE_PERSON', 'Property_Crime_Count', 'Person_Crime_Count',
    'G01_Population_Total',
    'G02_Median_age_persons', 'G02_Median_mortgage_repay_monthly',
    'G02_Median_tot_prsnl_inc_weekly', 'G02_Median_rent_weekly',
    'G02_Median_tot_fam_inc_weekly', 'G02_Average_num_psns_per_bedroom',
    'G02_Median_tot_hhd_inc_weekly', 'G02_Average_household_size',
]

_RENTAL_COLS = [
    'Suburb', 'Census_Rent_2021', 'Fair_Rent_2025', 'Fair_House_Rent_2025',
    'Fair_Unit_Rent_2025', 'Estimated_Actual_Rent_2025', 'Actual_House_Rent_2025',
    'Actual_Unit_Rent_2025', 'Greediness_Percent', 'Fair_House_Yield',
    'Actual_House_Yield', 'Actual_Unit_Yield', 'Affordability_Category',
    'Affordability_Ratio', 'Individual_Affordability', 'Household_Affordability',
]

_CULTURE_COLS = [
    'Suburb', 'G01_Population_Total', 'Current_Price_2025',
    'Indian_Population', 'Indian_Percent', 'Chinese_Population', 'Chinese_Percent',
    'Vietnamese_Population', 'Vietnamese_Percent', 'Italian_Population', 'Italian_Percent',
    'Greek_Population', 'Greek_Percent', 'Cultural_Diversity_Index',
]


def _safe_read_csv(path, usecols=None):
    """Read CSV with optional column selection, handling missing columns gracefully."""
    try:
        if usecols is not None:
            # Read header first to check which columns exist
            header = pd.read_csv(path, nrows=0).columns.tolist()
            valid_cols = [c for c in usecols if c in header]
            if not valid_cols:
                return pd.DataFrame()
            return pd.read_csv(path, usecols=valid_cols)
        return pd.read_csv(path)
    except FileNotFoundError:
        return pd.DataFrame()


@st.cache_data
def load_data():
    """Load and merge all data sources with optimized column selection."""
    # Master dataset — only load columns we need
    master_path = os.path.join(BASE, 'data/clean/master_dataset_by_suburb.csv')
    try:
        master_header = pd.read_csv(master_path, nrows=0).columns.tolist()
        valid_master_cols = [c for c in _MASTER_CORE_COLS if c in master_header]
        master = pd.read_csv(master_path, usecols=valid_master_cols)
    except FileNotFoundError:
        st.error("Master dataset not found!")
        return pd.DataFrame(), pd.DataFrame()

    df = master.copy()

    # Predictions — small file, load all
    preds = _safe_read_csv(os.path.join(BASE, 'data/predictions/price_predictions_2025_2026.csv'))
    if not preds.empty:
        pred_cols = [c for c in preds.columns if c not in df.columns or c == 'Suburb']
        df = df.merge(preds[pred_cols], on='Suburb', how='left')

    # Risk analysis — small file, load all
    risk = _safe_read_csv(os.path.join(BASE, 'data/risk_analysis/complete_risk_analysis.csv'))
    if not risk.empty:
        risk_cols = [c for c in risk.columns if c not in df.columns or c == 'Suburb']
        df = df.merge(risk[risk_cols], on='Suburb', how='left')

    # Rental analysis — HUGE file, only load needed columns
    rental = _safe_read_csv(
        os.path.join(BASE, 'data/rental/complete_rental_analysis.csv'),
        usecols=_RENTAL_COLS
    )
    if not rental.empty:
        rental_merge_cols = [c for c in rental.columns if c not in df.columns or c == 'Suburb']
        df = df.merge(rental[rental_merge_cols], on='Suburb', how='left')

    # Cultural demographics — selective load
    culture = _safe_read_csv(
        os.path.join(BASE, 'data/demographics/cultural_demographics.csv'),
        usecols=_CULTURE_COLS
    )
    if not culture.empty:
        culture_merge_cols = [c for c in culture.columns if c not in df.columns or c == 'Suburb']
        df = df.merge(culture[culture_merge_cols], on='Suburb', how='left')

    # Crime offense analysis — small file
    crime_off = _safe_read_csv(os.path.join(BASE, 'data/suburb_crime_offense_analysis.csv'))
    if not crime_off.empty:
        off_cols = [c for c in crime_off.columns if c not in df.columns or c == 'Suburb']
        df = df.merge(crime_off[off_cols], on='Suburb', how='left')

    # Timeseries — only load needed columns
    ts = _safe_read_csv(
        os.path.join(BASE, 'data/clean/property_timeseries_2019_2025.csv'),
        usecols=['Suburb', 'Median_Price', 'Period', 'Quarter', 'Year']
    )

    # Derived: Crime rate (only if not already present from risk analysis)
    if 'Crime_Rate_Per_1000' not in df.columns:
        if 'Total_Crime_Count' in df.columns and 'G01_Population_Total' in df.columns:
            df['Crime_Rate_Per_1000'] = np.where(
                df['G01_Population_Total'] > 0,
                df['Total_Crime_Count'] / df['G01_Population_Total'] * 1000, 0
            )

    return df, ts


@st.cache_data
def load_coordinates():
    """Load pre-built suburb coordinates."""
    coords_path = os.path.join(BASE, 'suburb_coordinates.json')
    if os.path.exists(coords_path):
        with open(coords_path, 'r') as f:
            return json.load(f)
    return {}


@st.cache_data
def load_geojson():
    """Load Adelaide suburb boundary GeoJSON."""
    geo_path = os.path.join(BASE, 'adelaide_suburbs.geojson')
    if os.path.exists(geo_path):
        with open(geo_path, 'r') as f:
            return json.load(f)
    return None


# ─────────────────────────────────────────────────────────────────────────────
# HELPER FUNCTIONS
# ─────────────────────────────────────────────────────────────────────────────

def fmt_price(val):
    """Format price as $XXX,XXX or $X.XXM."""
    if pd.isna(val):
        return "N/A"
    if val >= 1_000_000:
        return f"${val/1_000_000:.2f}M"
    return f"${val:,.0f}"


def fmt_dollar(val):
    """Format a dollar value — returns 'N/A' without $ prefix if missing."""
    if pd.isna(val):
        return "N/A"
    return f"${val:,.0f}"


def fmt_pct(val, decimals=1):
    if pd.isna(val):
        return "N/A"
    return f"{val:.{decimals}f}%"


def fmt_num(val, decimals=0):
    if pd.isna(val):
        return "N/A"
    return f"{val:,.{decimals}f}"


def get_val(row, col, default=np.nan):
    return row.get(col, default) if col in row.index else default


def price_tier_label(price):
    if pd.isna(price):
        return "Unknown"
    if price < 500_000:
        return "Budget (<$500K)"
    if price < 750_000:
        return "Mid-Range ($500K-$750K)"
    if price < 1_000_000:
        return "Upper-Mid ($750K-$1M)"
    return "Premium (>$1M)"


def price_tier_color(price):
    if pd.isna(price):
        return "#888888"
    if price < 500_000:
        return "#69f0ae"
    if price < 750_000:
        return "#ffd740"
    if price < 1_000_000:
        return "#ff6e40"
    return "#ff4081"


def risk_badge(category):
    if pd.isna(category):
        return ""
    cat = str(category).lower()
    if 'low' in cat:
        return f'<span class="badge badge-green">{category}</span>'
    if 'medium' in cat or 'moderate' in cat:
        return f'<span class="badge badge-orange">{category}</span>'
    return f'<span class="badge badge-red">{category}</span>'


def metric_card(label, value, delta=None, delta_type="neu"):
    delta_html = ""
    if delta:
        delta_html = f'<div class="delta delta-{delta_type}">{delta}</div>'
    return f"""
    <div class="metric-card">
        <div class="label">{label}</div>
        <div class="value">{value}</div>
        {delta_html}
    </div>
    """


# ─────────────────────────────────────────────────────────────────────────────
# PDF GENERATION
# ─────────────────────────────────────────────────────────────────────────────

def generate_pdf(suburb_name, row):
    """Generate a PDF report for a suburb."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)

    # Title
    pdf.set_font("Helvetica", "B", 20)
    pdf.set_text_color(30, 30, 80)
    pdf.cell(0, 15, f"{suburb_name} - Property Report", ln=True, align="C")
    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(100, 100, 100)
    pdf.cell(0, 8, "Adelaide Property Market Explorer | Data: 2019-2025", ln=True, align="C")
    pdf.ln(5)

    def section(title):
        pdf.set_font("Helvetica", "B", 13)
        pdf.set_text_color(50, 50, 120)
        pdf.cell(0, 10, title, ln=True)
        pdf.set_draw_color(100, 100, 200)
        pdf.line(10, pdf.get_y(), 200, pdf.get_y())
        pdf.ln(3)

    def field(label, value):
        pdf.set_font("Helvetica", "B", 10)
        pdf.set_text_color(60, 60, 60)
        pdf.cell(80, 7, label, ln=False)
        pdf.set_font("Helvetica", "", 10)
        pdf.set_text_color(30, 30, 30)
        pdf.cell(0, 7, str(value), ln=True)

    # Price Overview
    section("Price Overview")
    field("Current Price (2025):", fmt_price(get_val(row, 'Current_Price_2025')))
    field("First Price (2019):", fmt_price(get_val(row, 'First_Price_2019')))
    field("7-Year Growth:", fmt_pct(get_val(row, 'Price_Growth_Percent')))
    field("Growth Amount:", fmt_price(get_val(row, 'Price_Growth_Amount')))
    field("Price Volatility:", fmt_num(get_val(row, 'Price_Volatility'), 2))
    field("Price Tier:", price_tier_label(get_val(row, 'Current_Price_2025')))
    pdf.ln(3)

    # Demographics
    section("Demographics")
    field("Population:", fmt_num(get_val(row, 'G01_Population_Total')))
    field("Median Age:", fmt_num(get_val(row, 'G02_Median_age_persons'), 1))
    field("Median HH Income (weekly):", fmt_dollar(get_val(row, 'G02_Median_tot_hhd_inc_weekly')))
    field("Median Personal Income (weekly):", fmt_dollar(get_val(row, 'G02_Median_tot_prsnl_inc_weekly')))
    field("Median Mortgage Repay (monthly):", fmt_dollar(get_val(row, 'G02_Median_mortgage_repay_monthly')))
    pdf.ln(3)

    # Crime & Safety
    section("Crime & Safety")
    field("Total Crime Count:", fmt_num(get_val(row, 'Total_Crime_Count')))
    field("Property Crimes:", fmt_num(get_val(row, 'Property_Crime_Count')))
    field("Person Crimes:", fmt_num(get_val(row, 'Person_Crime_Count')))
    crime_rate = get_val(row, 'Crime_Rate_Per_1000')
    field("Crime Rate (per 1,000):", fmt_num(crime_rate, 1))
    pdf.ln(3)

    # Rental & Yield
    section("Rental & Yield")
    field("Fair House Rent 2025 (weekly):", fmt_dollar(get_val(row, 'Fair_House_Rent_2025')))
    field("Actual House Rent 2025 (weekly):", fmt_dollar(get_val(row, 'Actual_House_Rent_2025')))
    greed = get_val(row, 'Greediness_Percent')
    field("Greediness Gap:", fmt_pct(greed))
    field("Actual House Yield:", fmt_pct(get_val(row, 'Actual_House_Yield')))
    field("Affordability Category:", str(get_val(row, 'Affordability_Category', 'N/A')))
    pdf.ln(1)
    pdf.set_font("Helvetica", "I", 8)
    pdf.set_text_color(180, 100, 0)
    pdf.multi_cell(0, 5,
        "Disclaimer: Rental data is based on Census 2021, adjusted for inflation. "
        "For current rental prices, please check domain.com.au or realestate.com.au")
    pdf.ln(3)

    # Predictions & Risk
    section("Predictions & Risk")
    field("ML Predicted Price 2025:", fmt_price(get_val(row, 'Predicted_Price_2025')))
    field("Forecast Price 2026:", fmt_price(get_val(row, 'Forecast_Price_2026')))
    field("Expected Growth 2026:", fmt_pct(get_val(row, 'Expected_Growth_2026')))
    field("Total Risk Score:", fmt_num(get_val(row, 'Total_Risk_Score'), 1))
    field("Risk Category:", str(get_val(row, 'Total_Risk_Category', 'N/A')))
    field("Investment Strategy:", str(get_val(row, 'Investment_Strategy', 'N/A')))
    pdf.ln(3)

    # Crime Offense Breakdown
    ct1 = get_val(row, 'Crime_Type_1', '')
    if ct1 and str(ct1) != 'nan':
        section("Crime Offense Breakdown")
        field("Most Common:", str(ct1))
        ct2 = get_val(row, 'Crime_Type_2', '')
        if ct2 and str(ct2) != 'nan':
            field("2nd Most Common:", str(ct2))
        ct3 = get_val(row, 'Crime_Type_3', '')
        if ct3 and str(ct3) != 'nan':
            field("3rd Most Common:", str(ct3))
        pdf.ln(3)

    # Cultural
    section("Cultural Communities")
    for community in ['Indian', 'Chinese', 'Vietnamese', 'Italian', 'Greek']:
        pct = get_val(row, f'{community}_Percent')
        pop = get_val(row, f'{community}_Population')
        if not pd.isna(pct):
            field(f"{community}:", f"{fmt_pct(pct)} ({fmt_num(pop)} people)")
    div_idx = get_val(row, 'Cultural_Diversity_Index')
    field("Cultural Diversity Index:", fmt_num(div_idx, 3))

    # Footer
    pdf.ln(10)
    pdf.set_font("Helvetica", "I", 8)
    pdf.set_text_color(150, 150, 150)
    pdf.cell(0, 5, "Generated by Adelaide Property Market Explorer | Data: ABS Census 2021, SA Govt Crime Data, Property Records 2019-2025", ln=True, align="C")

    return bytes(pdf.output())


# ─────────────────────────────────────────────────────────────────────────────
# DOCX GENERATION
# ─────────────────────────────────────────────────────────────────────────────

def generate_docx(suburb_name, row):
    """Generate a DOCX report for a suburb."""
    doc = Document()

    # Styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    # Title
    title = doc.add_heading(f'{suburb_name} — Property Report', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph('Adelaide Property Market Explorer | Data: 2019–2025')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.color.rgb = RGBColor(120, 120, 120)
    p.runs[0].font.size = Pt(10)

    def section(title):
        doc.add_heading(title, level=2)

    def field(label, value):
        p = doc.add_paragraph()
        run = p.add_run(f'{label}  ')
        run.bold = True
        run.font.size = Pt(10)
        run2 = p.add_run(str(value))
        run2.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(2)

    # Price Overview
    section('Price Overview')
    field('Current Price (2025):', fmt_price(get_val(row, 'Current_Price_2025')))
    field('First Price (2019):', fmt_price(get_val(row, 'First_Price_2019')))
    field('7-Year Growth:', fmt_pct(get_val(row, 'Price_Growth_Percent')))
    field('Growth Amount:', fmt_price(get_val(row, 'Price_Growth_Amount')))
    field('Price Volatility:', fmt_num(get_val(row, 'Price_Volatility'), 2))
    field('Price Tier:', price_tier_label(get_val(row, 'Current_Price_2025')))

    # Demographics
    section('Demographics')
    field('Population:', fmt_num(get_val(row, 'G01_Population_Total')))
    field('Median Age:', fmt_num(get_val(row, 'G02_Median_age_persons'), 1))
    field('Median HH Income (weekly):', fmt_dollar(get_val(row, 'G02_Median_tot_hhd_inc_weekly')))
    field('Median Personal Income (weekly):', fmt_dollar(get_val(row, 'G02_Median_tot_prsnl_inc_weekly')))
    field('Median Mortgage (monthly):', fmt_dollar(get_val(row, 'G02_Median_mortgage_repay_monthly')))

    # Crime
    section('Crime & Safety')
    field('Total Crime Count:', fmt_num(get_val(row, 'Total_Crime_Count')))
    field('Property Crimes:', fmt_num(get_val(row, 'Property_Crime_Count')))
    field('Person Crimes:', fmt_num(get_val(row, 'Person_Crime_Count')))
    field('Crime Rate (per 1,000):', fmt_num(get_val(row, 'Crime_Rate_Per_1000'), 1))

    # Rental
    section('Rental & Yield')
    field('Fair House Rent 2025 (weekly):', fmt_dollar(get_val(row, 'Fair_House_Rent_2025')))
    field('Actual House Rent 2025 (weekly):', fmt_dollar(get_val(row, 'Actual_House_Rent_2025')))
    field('Greediness Gap:', fmt_pct(get_val(row, 'Greediness_Percent')))
    field('Actual House Yield:', fmt_pct(get_val(row, 'Actual_House_Yield')))
    field('Affordability Category:', str(get_val(row, 'Affordability_Category', 'N/A')))

    # Disclaimer
    p = doc.add_paragraph()
    run = p.add_run('Disclaimer: ')
    run.bold = True
    run.font.color.rgb = RGBColor(200, 120, 0)
    run.font.size = Pt(9)
    run2 = p.add_run(
        'Rental data is based on Census 2021, adjusted for inflation. '
        'For current rental prices, please check domain.com.au or realestate.com.au'
    )
    run2.font.color.rgb = RGBColor(150, 100, 0)
    run2.font.size = Pt(9)

    # Predictions
    section('Predictions & Risk')
    field('ML Predicted Price 2025:', fmt_price(get_val(row, 'Predicted_Price_2025')))
    field('Forecast Price 2026:', fmt_price(get_val(row, 'Forecast_Price_2026')))
    field('Expected Growth 2026:', fmt_pct(get_val(row, 'Expected_Growth_2026')))
    field('Total Risk Score:', fmt_num(get_val(row, 'Total_Risk_Score'), 1))
    field('Risk Category:', str(get_val(row, 'Total_Risk_Category', 'N/A')))
    field('Investment Strategy:', str(get_val(row, 'Investment_Strategy', 'N/A')))

    # Cultural
    section('Cultural Communities')
    for community in ['Indian', 'Chinese', 'Vietnamese', 'Italian', 'Greek']:
        pct = get_val(row, f'{community}_Percent')
        pop = get_val(row, f'{community}_Population')
        if not pd.isna(pct):
            field(f'{community}:', f"{fmt_pct(pct)} ({fmt_num(pop)} people)")
    field('Cultural Diversity Index:', fmt_num(get_val(row, 'Cultural_Diversity_Index'), 3))

    # Crime Offense Breakdown
    ct1 = get_val(row, 'Crime_Type_1', '')
    if ct1 and str(ct1) != 'nan':
        section('Crime Offense Breakdown')
        field('Most Common:', str(ct1))
        ct2 = get_val(row, 'Crime_Type_2', '')
        if ct2 and str(ct2) != 'nan':
            field('2nd Most Common:', str(ct2))
        ct3 = get_val(row, 'Crime_Type_3', '')
        if ct3 and str(ct3) != 'nan':
            field('3rd Most Common:', str(ct3))

    # Footer
    doc.add_paragraph()
    p = doc.add_paragraph('Generated by Adelaide Property Market Explorer | Data: ABS Census 2021, SA Govt, Property Records 2019-2025')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.color.rgb = RGBColor(160, 160, 160)
    p.runs[0].font.size = Pt(8)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ─────────────────────────────────────────────────────────────────────────────
# PRICE HISTORY CHART — Now using Plotly for client-side rendering
# ─────────────────────────────────────────────────────────────────────────────

def create_price_chart(ts, suburb_name):
    """Create an interactive Plotly price history chart for the suburb."""
    suburb_ts = ts[ts['Suburb'] == suburb_name].copy()
    if suburb_ts.empty:
        return None

    # Safe parsing of Period column
    year_extract = suburb_ts['Period'].str.extract(r'(\d{4})')
    qnum_extract = suburb_ts['Period'].str.extract(r'Q(\d)')

    if year_extract[0].isna().all() or qnum_extract[0].isna().all():
        return None

    suburb_ts['_year'] = pd.to_numeric(year_extract[0], errors='coerce')
    suburb_ts['_qnum'] = pd.to_numeric(qnum_extract[0], errors='coerce')
    suburb_ts = suburb_ts.dropna(subset=['_year', '_qnum'])
    suburb_ts['_year'] = suburb_ts['_year'].astype(int)
    suburb_ts['_qnum'] = suburb_ts['_qnum'].astype(int)
    suburb_ts = suburb_ts.sort_values(['_year', '_qnum']).reset_index(drop=True)
    suburb_ts['Label'] = suburb_ts['_year'].astype(str) + ' Q' + suburb_ts['_qnum'].astype(str)

    fig = go.Figure()

    fig.add_trace(go.Scatter(
        x=suburb_ts['Label'],
        y=suburb_ts['Median_Price'],
        mode='lines+markers',
        name='Median Price',
        line=dict(color='#667eea', width=3),
        marker=dict(size=7, color='#a78bfa', line=dict(width=2, color='#667eea')),
        fill='tozeroy',
        fillcolor='rgba(102, 126, 234, 0.1)',
        hovertemplate='<b>%{x}</b><br>Price: $%{y:,.0f}<extra></extra>',
    ))

    low = suburb_ts['Median_Price'].min()
    high = suburb_ts['Median_Price'].max()

    fig.update_layout(
        title=dict(
            text=f'{suburb_name} — Price History ({fmt_price(low)} → {fmt_price(high)})',
            font=dict(size=16, color='#e0e0ff'),
        ),
        plot_bgcolor='#1a1a2e',
        paper_bgcolor='#0e1117',
        font=dict(color='#b8b8d0'),
        xaxis=dict(
            showgrid=True, gridcolor='rgba(74,74,106,0.3)',
            tickangle=-45, dtick=4,
        ),
        yaxis=dict(
            showgrid=True, gridcolor='rgba(74,74,106,0.3)',
            tickformat='$,.0f',
            title='Median Price',
        ),
        hovermode='x unified',
        height=400,
        margin=dict(l=60, r=20, t=50, b=80),
    )

    return fig


# ─────────────────────────────────────────────────────────────────────────────
# YEAR-TO-YEAR GROWTH ANALYSIS
# ─────────────────────────────────────────────────────────────────────────────

def compute_yoy_growth(ts, suburb_name):
    """Compute year-over-year median price and growth for a suburb."""
    suburb_ts = ts[ts['Suburb'] == suburb_name].copy()
    if suburb_ts.empty:
        return None

    # Parse year from Period
    year_extract = suburb_ts['Period'].str.extract(r'(\d{4})')
    suburb_ts['_year'] = pd.to_numeric(year_extract[0], errors='coerce')
    suburb_ts = suburb_ts.dropna(subset=['_year', 'Median_Price'])
    suburb_ts['_year'] = suburb_ts['_year'].astype(int)

    # Annual median price (average of all quarters in each year)
    annual = suburb_ts.groupby('_year')['Median_Price'].median().reset_index()
    annual.columns = ['Year', 'Median_Price']
    annual = annual.sort_values('Year').reset_index(drop=True)

    # Year-over-year growth
    annual['YoY_Growth_Amount'] = annual['Median_Price'].diff()
    annual['YoY_Growth_Pct'] = annual['Median_Price'].pct_change() * 100

    # Guard against inf values from zero prior prices
    annual['YoY_Growth_Pct'] = annual['YoY_Growth_Pct'].replace([np.inf, -np.inf], np.nan)

    return annual


def create_yoy_chart(annual_data, suburb_name):
    """Create a combined bar + line chart for YoY growth."""
    if annual_data is None or annual_data.empty:
        return None

    fig = go.Figure()

    # Bar chart for annual median price
    colors = ['#69f0ae' if g >= 0 else '#ff5252'
              for g in annual_data['YoY_Growth_Pct'].fillna(0)]

    fig.add_trace(go.Bar(
        x=annual_data['Year'].astype(str),
        y=annual_data['Median_Price'],
        name='Median Price',
        marker_color='rgba(102, 126, 234, 0.6)',
        marker_line=dict(color='#667eea', width=1.5),
        hovertemplate='<b>%{x}</b><br>Price: $%{y:,.0f}<extra></extra>',
    ))

    # Line chart for YoY growth % on secondary axis
    fig.add_trace(go.Scatter(
        x=annual_data['Year'].astype(str),
        y=annual_data['YoY_Growth_Pct'],
        name='YoY Growth %',
        mode='lines+markers+text',
        line=dict(color='#f093fb', width=3),
        marker=dict(size=10, color=colors, line=dict(width=2, color='#f093fb')),
        text=[f"{v:+.1f}%" if not pd.isna(v) else "" for v in annual_data['YoY_Growth_Pct']],
        textposition='top center',
        textfont=dict(size=11, color='#e0e0ff'),
        yaxis='y2',
        hovertemplate='<b>%{x}</b><br>Growth: %{y:.1f}%<extra></extra>',
    ))

    fig.update_layout(
        title=dict(
            text=f'{suburb_name} — Year-to-Year Price Growth',
            font=dict(size=16, color='#e0e0ff'),
        ),
        plot_bgcolor='#1a1a2e',
        paper_bgcolor='#0e1117',
        font=dict(color='#b8b8d0'),
        xaxis=dict(showgrid=False, title='Year'),
        yaxis=dict(
            showgrid=True, gridcolor='rgba(74,74,106,0.3)',
            tickformat='$,.0f', title='Median Price',
        ),
        yaxis2=dict(
            title='YoY Growth %', overlaying='y', side='right',
            showgrid=False, tickformat='+.1f',
            zeroline=True, zerolinecolor='rgba(255,255,255,0.3)',
        ),
        barmode='group',
        hovermode='x unified',
        height=450,
        margin=dict(l=60, r=60, t=50, b=50),
        legend=dict(
            orientation='h', yanchor='bottom', y=1.02,
            xanchor='center', x=0.5,
            font=dict(color='#c8c8e0'),
        ),
    )

    return fig


def render_yoy_tab(ts, suburb_name):
    """Render the Year-to-Year Growth tab content."""
    annual = compute_yoy_growth(ts, suburb_name)

    if annual is None or annual.empty:
        st.info("No timeseries data available for year-to-year analysis.")
        return

    # YoY Chart
    fig = create_yoy_chart(annual, suburb_name)
    if fig:
        st.plotly_chart(fig, use_container_width=True)

    # Summary cards
    st.markdown('<div class="section-header">Annual Breakdown</div>', unsafe_allow_html=True)

    years = annual.sort_values('Year', ascending=False)
    for _, yr_row in years.iterrows():
        year = int(yr_row['Year'])
        price = yr_row['Median_Price']
        growth_pct = yr_row['YoY_Growth_Pct']
        growth_amt = yr_row['YoY_Growth_Amount']

        if pd.isna(growth_pct):
            growth_html = '<span style="color: #9999bb;">— Base Year —</span>'
        elif growth_pct >= 0:
            growth_html = (f'<span class="growth-up">+{growth_pct:.1f}% '
                           f'(+{fmt_price(growth_amt)})</span>')
        else:
            growth_html = (f'<span class="growth-down">{growth_pct:.1f}% '
                           f'({fmt_price(growth_amt)})</span>')

        st.markdown(f"""
        <div class="yoy-card">
            <span class="year">{year}</span>
            &nbsp;&nbsp;
            <span class="price">{fmt_price(price)}</span>
            &nbsp;&nbsp;
            {growth_html}
        </div>
        """, unsafe_allow_html=True)

    # CAGR calculation
    if len(annual) >= 2:
        first_price = annual.iloc[0]['Median_Price']
        last_price = annual.iloc[-1]['Median_Price']
        n_years = annual.iloc[-1]['Year'] - annual.iloc[0]['Year']
        if first_price > 0 and n_years > 0:
            cagr = ((last_price / first_price) ** (1 / n_years) - 1) * 100
            total_growth = ((last_price - first_price) / first_price) * 100

            st.markdown('<div class="section-header">Growth Summary</div>',
                        unsafe_allow_html=True)
            c1, c2, c3, c4 = st.columns(4)
            with c1:
                st.markdown(metric_card("CAGR", fmt_pct(cagr),
                                        f"{n_years} years",
                                        "pos" if cagr > 0 else "neg"),
                            unsafe_allow_html=True)
            with c2:
                st.markdown(metric_card("Total Growth", fmt_pct(total_growth),
                                        f"{int(annual.iloc[0]['Year'])}-{int(annual.iloc[-1]['Year'])}",
                                        "pos" if total_growth > 0 else "neg"),
                            unsafe_allow_html=True)
            with c3:
                best_yr = annual.loc[annual['YoY_Growth_Pct'].idxmax()] if annual['YoY_Growth_Pct'].notna().any() else None
                if best_yr is not None:
                    st.markdown(metric_card("Best Year",
                                            str(int(best_yr['Year'])),
                                            f"+{best_yr['YoY_Growth_Pct']:.1f}%", "pos"),
                                unsafe_allow_html=True)
            with c4:
                worst_yr = annual.loc[annual['YoY_Growth_Pct'].idxmin()] if annual['YoY_Growth_Pct'].notna().any() else None
                if worst_yr is not None:
                    dtype = "neg" if worst_yr['YoY_Growth_Pct'] < 0 else "neu"
                    st.markdown(metric_card("Worst Year",
                                            str(int(worst_yr['Year'])),
                                            f"{worst_yr['YoY_Growth_Pct']:.1f}%", dtype),
                                unsafe_allow_html=True)


# ─────────────────────────────────────────────────────────────────────────────
# MAP
# ─────────────────────────────────────────────────────────────────────────────

def create_map(df, coords, geojson_data=None):
    """Create an interactive Adelaide map — single GeoJson layer for speed."""
    m = folium.Map(location=[-34.9285, 138.6007], zoom_start=11,
                   tiles="CartoDB dark_matter")

    # Build data lookups
    data_lookup = {}
    for _, row in df.iterrows():
        suburb = row['Suburb']
        price = row.get('Current_Price_2025', np.nan)
        data_lookup[suburb] = {
            'price': price,
            'price_fmt': fmt_price(price),
            'growth': row.get('Price_Growth_Percent', np.nan),
            'growth_fmt': fmt_pct(row.get('Price_Growth_Percent', np.nan)),
            'pred_2026': row.get('Forecast_Price_2026', np.nan),
            'pred_fmt': fmt_price(row.get('Forecast_Price_2026', np.nan)),
            'exp_growth': fmt_pct(row.get('Expected_Growth_2026', np.nan)),
            'color': price_tier_color(price),
        }

    if geojson_data:
        for feature in geojson_data['features']:
            suburb = feature['properties'].get('Suburb', '')
            info = data_lookup.get(suburb, {})
            feature['properties']['price'] = info.get('price_fmt', 'N/A')
            feature['properties']['growth'] = info.get('growth_fmt', 'N/A')
            feature['properties']['pred_2026'] = info.get('pred_fmt', 'N/A')
            feature['properties']['exp_growth'] = info.get('exp_growth', 'N/A')
            feature['properties']['fill_color'] = info.get('color', '#888888')

        folium.GeoJson(
            geojson_data,
            style_function=lambda feature: {
                'fillColor': feature['properties'].get('fill_color', '#888888'),
                'color': '#ffffff',
                'weight': 0.6,
                'fillOpacity': 0.55,
            },
            highlight_function=lambda feature: {
                'weight': 3,
                'color': '#ffffff',
                'fillOpacity': 0.85,
            },
            tooltip=folium.GeoJsonTooltip(
                fields=['Suburb', 'price', 'growth', 'pred_2026', 'exp_growth'],
                aliases=['Suburb:', 'Current Price:', '7-Year Growth:', 'Predicted 2026:', 'Forecast Growth:'],
                style='background-color: rgba(15,12,41,0.9); color: #e0e0ff; '
                      'font-family: Inter, sans-serif; font-size: 12px; '
                      'padding: 8px; border-radius: 6px; border: 1px solid #4a4a8a;',
                sticky=True,
            ),
        ).add_to(m)
    else:
        # Fallback: circle markers
        for _, row in df.iterrows():
            suburb = row['Suburb']
            if suburb not in coords:
                continue
            coord = coords[suburb]
            price = row.get('Current_Price_2025', np.nan)
            color = price_tier_color(price)
            folium.CircleMarker(
                location=[coord['lat'], coord['lng']],
                radius=6, color=color, fill=True,
                fill_color=color, fill_opacity=0.7,
                tooltip=f"{suburb}: {fmt_price(price)}"
            ).add_to(m)

    # Legend
    legend_html = """
    <div style="position: fixed; bottom: 30px; left: 30px; z-index: 1000;
                background: rgba(15,12,41,0.92); padding: 14px 18px;
                border-radius: 12px; border: 1px solid #4a4a8a;
                font-family: Inter, sans-serif; font-size: 12px; color: #e0e0ff;
                box-shadow: 0 4px 20px rgba(0,0,0,0.4);">
        <b style="color: #f093fb;">Price Tier</b><br>
        <span style="color: #69f0ae;">&#9679;</span> Budget (&lt;$500K)<br>
        <span style="color: #ffd740;">&#9679;</span> Mid-Range ($500K–$750K)<br>
        <span style="color: #ff6e40;">&#9679;</span> Upper-Mid ($750K–$1M)<br>
        <span style="color: #ff4081;">&#9679;</span> Premium (&gt;$1M)
    </div>
    """
    m.get_root().html.add_child(folium.Element(legend_html))

    return m


# ─────────────────────────────────────────────────────────────────────────────
# SUBURB REPORT — Now with tabs
# ─────────────────────────────────────────────────────────────────────────────

def render_suburb_report(row, ts):
    """Render the full suburb report with tabbed navigation."""
    suburb = row['Suburb']

    st.markdown(f"""
    <div style="text-align: center; margin: 1rem 0;">
        <span style="font-size: 2rem; font-weight: 800;
                     background: linear-gradient(90deg, #667eea, #764ba2, #f093fb);
                     -webkit-background-clip: text; -webkit-text-fill-color: transparent;">
            {suburb}
        </span>
    </div>
    """, unsafe_allow_html=True)

    # Tabbed navigation
    tab_overview, tab_yoy, tab_demo, tab_crime, tab_rental, tab_predict, tab_culture, tab_download = st.tabs([
        "Overview", "YoY Growth", "Demographics",
        "Crime & Safety", "Rental & Yield", "Predictions & Risk",
        "Cultural", "Download"
    ])

    # ── TAB: OVERVIEW ──
    with tab_overview:
        # Show a note if this suburb has missing price data
        current_price = get_val(row, 'Current_Price_2025')
        if pd.isna(current_price):
            st.warning(
                f"**{suburb}** has limited property data. "
                f"This suburb may have too few property sales in recent quarters to calculate "
                f"reliable statistics. 44 of 414 suburbs lack 2025 price data, which also means "
                f"predictions and risk analysis are unavailable."
            )

        st.markdown(f'<div class="section-header">Price Overview</div>',
                    unsafe_allow_html=True)

        c1, c2, c3, c4 = st.columns(4)
        with c1:
            st.markdown(metric_card("Current Price 2025",
                                    fmt_price(get_val(row, 'Current_Price_2025'))),
                        unsafe_allow_html=True)
        with c2:
            growth = get_val(row, 'Price_Growth_Percent')
            delta_type = "pos" if not pd.isna(growth) and growth > 0 else "neg"
            st.markdown(metric_card("7-Year Growth", fmt_pct(growth),
                                    "Since 2019", delta_type),
                        unsafe_allow_html=True)
        with c3:
            st.markdown(metric_card("First Price 2019",
                                    fmt_price(get_val(row, 'First_Price_2019'))),
                        unsafe_allow_html=True)
        with c4:
            tier = price_tier_label(get_val(row, 'Current_Price_2025'))
            st.markdown(metric_card("Market Tier", tier), unsafe_allow_html=True)

        # Price History Chart
        st.markdown('<div class="section-header">Price History (2019–2025)</div>',
                    unsafe_allow_html=True)
        fig = create_price_chart(ts, suburb)
        if fig:
            st.plotly_chart(fig, use_container_width=True)
        else:
            st.info("No timeseries data available for this suburb.")

        # Quick stats row
        c1, c2, c3 = st.columns(3)
        with c1:
            st.markdown(metric_card("Price Volatility",
                                    fmt_num(get_val(row, 'Price_Volatility'), 2),
                                    "Std deviation", "neu"),
                        unsafe_allow_html=True)
        with c2:
            st.markdown(metric_card("Growth Amount",
                                    fmt_price(get_val(row, 'Price_Growth_Amount')),
                                    "2019 → 2025", "pos"),
                        unsafe_allow_html=True)
        with c3:
            st.markdown(metric_card("Data Points",
                                    fmt_num(get_val(row, 'Quarter_Count')),
                                    "Quarterly records", "neu"),
                        unsafe_allow_html=True)

    # ── TAB: YEAR-TO-YEAR GROWTH ──
    with tab_yoy:
        st.markdown('<div class="section-header">Year-to-Year Growth Analysis</div>',
                    unsafe_allow_html=True)
        render_yoy_tab(ts, suburb)

    # ── TAB: DEMOGRAPHICS ──
    with tab_demo:
        st.markdown('<div class="section-header">Demographics</div>',
                    unsafe_allow_html=True)
        c1, c2, c3 = st.columns(3)
        with c1:
            st.markdown(metric_card("Population",
                                    fmt_num(get_val(row, 'G01_Population_Total'))),
                        unsafe_allow_html=True)
        with c2:
            st.markdown(metric_card("Median Age",
                                    fmt_num(get_val(row, 'G02_Median_age_persons'), 1)),
                        unsafe_allow_html=True)
        with c3:
            st.markdown(metric_card("Household Size",
                                    fmt_num(get_val(row, 'G02_Average_household_size'), 1)),
                        unsafe_allow_html=True)

        st.markdown("")  # spacer
        c1, c2, c3 = st.columns(3)
        with c1:
            st.markdown(metric_card("HH Income (wk)",
                                    fmt_dollar(get_val(row, 'G02_Median_tot_hhd_inc_weekly'))),
                        unsafe_allow_html=True)
        with c2:
            st.markdown(metric_card("Personal Income (wk)",
                                    fmt_dollar(get_val(row, 'G02_Median_tot_prsnl_inc_weekly'))),
                        unsafe_allow_html=True)
        with c3:
            st.markdown(metric_card("Mortgage (mth)",
                                    fmt_dollar(get_val(row, 'G02_Median_mortgage_repay_monthly'))),
                        unsafe_allow_html=True)

        st.markdown("""
        <div class="disclaimer-box">
            <strong>Census Data Note:</strong> All demographic figures (population, age, income,
            mortgage) are from the <strong>ABS Census 2021</strong> — the most recent census available.
            Actual figures may differ from 2021 values.
        </div>
        """, unsafe_allow_html=True)


    # ── TAB: CRIME & SAFETY ──
    with tab_crime:
        st.markdown('<div class="section-header">Crime & Safety</div>',
                    unsafe_allow_html=True)
        c1, c2, c3, c4 = st.columns(4)
        with c1:
            st.markdown(metric_card("Total Crimes",
                                    fmt_num(get_val(row, 'Total_Crime_Count'))),
                        unsafe_allow_html=True)
        with c2:
            st.markdown(metric_card("Property Crimes",
                                    fmt_num(get_val(row, 'Property_Crime_Count'))),
                        unsafe_allow_html=True)
        with c3:
            st.markdown(metric_card("Person Crimes",
                                    fmt_num(get_val(row, 'Person_Crime_Count'))),
                        unsafe_allow_html=True)
        with c4:
            crime_rate = get_val(row, 'Crime_Rate_Per_1000')
            delta_type = "pos" if not pd.isna(crime_rate) and crime_rate < 50 else \
                "neu" if not pd.isna(crime_rate) and crime_rate < 150 else "neg"
            st.markdown(metric_card("Crime / 1,000 ppl",
                                    fmt_num(crime_rate, 1),
                                    "Low" if delta_type == "pos" else "Moderate" if delta_type == "neu" else "High",
                                    delta_type),
                        unsafe_allow_html=True)

        # Crime Offense Breakdown
        ct1 = get_val(row, 'Crime_Type_1', '')
        ct2 = get_val(row, 'Crime_Type_2', '')
        ct3 = get_val(row, 'Crime_Type_3', '')
        no_of_crimes = get_val(row, 'No_of_Crimes')
        if ct1 and str(ct1) != 'nan':
            count_label = f" — Total Recorded: {fmt_num(no_of_crimes)}" if not pd.isna(no_of_crimes) else ""
            st.markdown(f'<div class="section-header">Top Offense Types{count_label}</div>',
                        unsafe_allow_html=True)
            oc1, oc2, oc3 = st.columns(3)
            with oc1:
                st.markdown(f'<div class="metric-card"><div class="label">#1 Most Common</div>'
                            f'<div class="value" style="font-size:1rem;">{ct1}</div></div>',
                            unsafe_allow_html=True)
            with oc2:
                if ct2 and str(ct2) != 'nan':
                    st.markdown(f'<div class="metric-card"><div class="label">#2 Most Common</div>'
                                f'<div class="value" style="font-size:1rem;">{ct2}</div></div>',
                                unsafe_allow_html=True)
            with oc3:
                if ct3 and str(ct3) != 'nan':
                    st.markdown(f'<div class="metric-card"><div class="label">#3 Most Common</div>'
                                f'<div class="value" style="font-size:1rem;">{ct3}</div></div>',
                                unsafe_allow_html=True)

        # Crime data period note
        st.markdown("""
        <div class="disclaimer-box">
            <strong>Crime Data Note:</strong> Crime statistics cover SA Government records from
            <strong>FY 2019-20 through Q1 2025-26</strong> (cumulative totals across all reporting periods).
        </div>
        """, unsafe_allow_html=True)

    # ── TAB: RENTAL & YIELD ──
    with tab_rental:
        st.markdown('<div class="section-header">Rental & Yield</div>',
                    unsafe_allow_html=True)

        has_rental = not pd.isna(get_val(row, 'Fair_House_Rent_2025'))
        if has_rental:
            c1, c2, c3 = st.columns(3)
            with c1:
                st.markdown(metric_card("Fair Rent (wk)",
                                        fmt_dollar(get_val(row, 'Fair_House_Rent_2025'))),
                            unsafe_allow_html=True)
            with c2:
                st.markdown(metric_card("Actual Rent (wk)",
                                        fmt_dollar(get_val(row, 'Actual_House_Rent_2025'))),
                            unsafe_allow_html=True)
            with c3:
                greed = get_val(row, 'Greediness_Percent')
                delta_type = "neg" if not pd.isna(greed) and greed > 10 else "neu"
                st.markdown(metric_card("Greediness Gap",
                                        fmt_pct(greed),
                                        "Overcharged" if delta_type == "neg" else "Fair",
                                        delta_type),
                            unsafe_allow_html=True)

            st.markdown("")
            c1, c2 = st.columns(2)
            with c1:
                st.markdown(metric_card("House Yield",
                                        fmt_pct(get_val(row, 'Actual_House_Yield'))),
                            unsafe_allow_html=True)
            with c2:
                afford = str(get_val(row, 'Affordability_Category', 'N/A'))
                st.markdown(metric_card("Affordability", afford), unsafe_allow_html=True)

            # Disclaimer
            st.markdown("""
            <div class="disclaimer-box">
                <strong>Rental Data Disclaimer:</strong> Rental figures are estimated from
                <strong>Census 2021</strong> data, adjusted for inflation (~30% increase 2021–2025).
                For current, accurate rental prices, please check
                <a href="https://www.domain.com.au" target="_blank" style="color: #ffa726;">domain.com.au</a> or
                <a href="https://www.realestate.com.au" target="_blank" style="color: #ffa726;">realestate.com.au</a>.
            </div>
            """, unsafe_allow_html=True)
        else:
            st.warning(
                f"**Rental data is not available for {suburb}.**\n\n"
                f"Only **368 of 414** suburbs have rental analysis (based on Census 2021 rental records). "
                f"This suburb may not have enough rental dwellings recorded in the census. "
                f"Check [domain.com.au](https://www.domain.com.au) or "
                f"[realestate.com.au](https://www.realestate.com.au) for current rental prices."
            )

    # ── TAB: PREDICTIONS & RISK ──
    with tab_predict:
        st.markdown('<div class="section-header">Predictions & Risk</div>',
                    unsafe_allow_html=True)
        has_predictions = not pd.isna(get_val(row, 'Predicted_Price_2025'))
        if has_predictions:
            c1, c2, c3 = st.columns(3)
            with c1:
                st.markdown(metric_card("ML Predicted 2025",
                                        fmt_price(get_val(row, 'Predicted_Price_2025'))),
                            unsafe_allow_html=True)
            with c2:
                exp_growth = get_val(row, 'Expected_Growth_2026')
                growth_delta_type = "pos" if not pd.isna(exp_growth) and exp_growth > 0 else "neg"
                st.markdown(metric_card("Forecast 2026",
                                        fmt_price(get_val(row, 'Forecast_Price_2026')),
                                        f"Growth: {fmt_pct(exp_growth)}",
                                        growth_delta_type),
                            unsafe_allow_html=True)
            with c3:
                risk_cat = str(get_val(row, 'Total_Risk_Category', 'N/A'))
                risk_score = get_val(row, 'Total_Risk_Score')
                st.markdown(metric_card("Risk Score",
                                        fmt_num(risk_score, 1),
                                        risk_cat,
                                        "pos" if 'Low' in risk_cat else "neg" if 'High' in risk_cat else "neu"),
                            unsafe_allow_html=True)

            st.markdown("")
            c1, c2, c3 = st.columns(3)
            with c1:
                strat = str(get_val(row, 'Investment_Strategy', 'N/A'))
                st.markdown(metric_card("Investment Strategy", strat), unsafe_allow_html=True)
            with c2:
                st.markdown(metric_card("Risk-Adj. Return",
                                        fmt_num(get_val(row, 'Risk_Adjusted_Return'), 2)),
                            unsafe_allow_html=True)
            with c3:
                st.markdown(metric_card("Value Category",
                                        str(get_val(row, 'Value_Category', 'N/A'))),
                            unsafe_allow_html=True)
        else:
            st.warning(
                f"**Prediction & risk data is not available for {suburb}.**\n\n"
                f"Only **370 of 414** suburbs have ML predictions. The remaining 44 suburbs lack "
                f"sufficient recent property sales data (no 2025 price) to generate reliable "
                f"price predictions or risk assessments."
            )

    # ── TAB: CULTURAL ──
    with tab_culture:
        st.markdown('<div class="section-header">Cultural Communities</div>',
                    unsafe_allow_html=True)

        communities = ['Indian', 'Chinese', 'Vietnamese', 'Italian', 'Greek']
        has_culture = not pd.isna(get_val(row, 'Indian_Percent'))
        if has_culture:
            c1, c2, c3 = st.columns(3)
            cols_cycle = [c1, c2, c3]
            for i, comm in enumerate(communities):
                pct = get_val(row, f'{comm}_Percent')
                pop = get_val(row, f'{comm}_Population')
                with cols_cycle[i % 3]:
                    subtitle = f"{fmt_num(pop)} people" if not pd.isna(pop) else ""
                    st.markdown(metric_card(comm, fmt_pct(pct), subtitle, "neu"),
                                unsafe_allow_html=True)

            st.markdown("")
            st.markdown(metric_card("Cultural Diversity Index",
                                    fmt_num(get_val(row, 'Cultural_Diversity_Index'), 3),
                                    "0 = homogeneous, 1 = diverse", "neu"),
                        unsafe_allow_html=True)

            st.markdown("""
            <div class="disclaimer-box">
                <strong>Demographics Note:</strong> Cultural demographics are sourced from
                <strong>ABS Census 2021</strong> (most recent census data available).
                Actual community composition may have shifted since 2021.
            </div>
            """, unsafe_allow_html=True)
        else:
            st.warning(
                f"**Cultural demographics data is not available for {suburb}.**\n\n"
                f"**412 of 414** suburbs have cultural demographics from Census 2021. "
                f"This suburb may not have been separately enumerated in the census."
            )

    # ── TAB: DOWNLOAD ──
    with tab_download:
        st.markdown('<div class="section-header">Download Report</div>',
                    unsafe_allow_html=True)
        st.markdown(f"Generate a comprehensive report for **{suburb}** including all "
                    f"price, demographic, crime, rental, prediction and cultural data.")

        c1, c2, _ = st.columns([1, 1, 3])
        with c1:
            if st.button("Generate PDF", key="gen_pdf"):
                with st.spinner("Generating PDF..."):
                    pdf_bytes = generate_pdf(suburb, row)
                st.download_button(
                    "Download PDF",
                    data=pdf_bytes,
                    file_name=f"{suburb}_property_report.pdf",
                    mime="application/pdf"
                )
        with c2:
            if st.button("Generate DOCX", key="gen_docx"):
                with st.spinner("Generating DOCX..."):
                    docx_bytes = generate_docx(suburb, row)
                st.download_button(
                    "Download DOCX",
                    data=docx_bytes,
                    file_name=f"{suburb}_property_report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )


# ─────────────────────────────────────────────────────────────────────────────
# MAIN APP
# ─────────────────────────────────────────────────────────────────────────────

def main():
    # Load data with spinner
    with st.spinner("Loading Adelaide property data..."):
        df, ts = load_data()
    coords = load_coordinates()
    geojson_data = load_geojson()

    if df.empty:
        st.error("Failed to load data. Please check the data directory.")
        return

    suburbs_list = sorted(df['Suburb'].unique())
    n_with_data = int(df['Current_Price_2025'].notna().sum())

    # Session state
    if 'map_suburb' not in st.session_state:
        st.session_state.map_suburb = None
    if 'last_click_id' not in st.session_state:
        st.session_state.last_click_id = None

    # Header
    st.markdown(f"""
    <div class="main-header">
        <h1>Adelaide Property Market Explorer</h1>
        <p>{n_with_data} Suburbs · 7 Years · Demographics · Crime · Growth · ML Predictions · Rental Yields</p>
    </div>
    """, unsafe_allow_html=True)

    # ── SIDEBAR ──
    with st.sidebar:
        st.markdown("## Find a Suburb")

        search_input = st.text_input("Enter suburb name:",
                                     placeholder="e.g. ADELAIDE, GLENELG, SALISBURY...")

        if search_input:
            search_upper = search_input.strip().upper()
            matches = [s for s in suburbs_list if search_upper in s]
            if not matches:
                st.warning(f"No suburb found matching '{search_input}'")
                matches = suburbs_list
        else:
            matches = suburbs_list

        default_idx = 0
        if st.session_state.map_suburb and st.session_state.map_suburb in matches:
            default_idx = matches.index(st.session_state.map_suburb)

        selected_suburb = st.selectbox("Select suburb:", matches, index=default_idx)

        if selected_suburb != st.session_state.map_suburb:
            st.session_state.map_suburb = None

        st.markdown("---")
        st.markdown("""
        **Data Sources & Periods:**
        - Property Prices — Q1 2019 to Q4 2025
        - SA Crime Statistics — FY 2019-20 to Q1 2025-26
        - ABS Census 2021 — Demographics & Rental
        - Rental Market — Census 2021, inflation-adjusted
        - Cultural Demographics — Census 2021
        """)
        st.markdown("""
        <div style="background: rgba(255,152,0,0.08); border-left: 3px solid #ff9800;
                    border-radius: 0 8px 8px 0; padding: 0.7rem 0.9rem; margin: 0.5rem 0;
                    font-size: 0.78rem; color: #ffcc80;">
            <strong style="color:#ffa726;">Data Disclaimer:</strong><br>
            Demographics & rental data are from <strong>ABS Census 2021</strong> (most recent available).
            Crime data covers SA Govt records through <strong>Q1 2025-26</strong>.
            Property prices span <strong>Q1 2019 – Q4 2025</strong> (28 quarters).
            Rental figures are inflation-adjusted estimates — check
            <a href="https://www.domain.com.au" target="_blank" style="color:#ffa726;">domain.com.au</a>
            for current rates.
        </div>
        """, unsafe_allow_html=True)
        st.markdown("---")
        st.markdown("**Created by** Abdul Mussavir")

    # ── LANDING DASHBOARD ──
    st.markdown('<div class="section-header">Market Snapshot — Top Performers & Risk Alerts</div>',
                unsafe_allow_html=True)

    valid = df[df['Current_Price_2025'].notna()].copy()

    col_left, col_right = st.columns(2)

    with col_left:
        st.markdown('#### Top 10 Growth Suburbs (7-Year)')
        if 'Price_Growth_Percent' in valid.columns:
            top_growth = valid.nlargest(10, 'Price_Growth_Percent')[[
                'Suburb', 'Current_Price_2025', 'Price_Growth_Percent'
            ]].copy()
            top_growth['Current_Price_2025'] = top_growth['Current_Price_2025'].apply(fmt_price)
            top_growth['Price_Growth_Percent'] = top_growth['Price_Growth_Percent'].apply(
                lambda x: fmt_pct(x))
            top_growth.columns = ['Suburb', 'Price 2025', 'Growth %']
            st.dataframe(top_growth, use_container_width=True, hide_index=True)

    with col_right:
        st.markdown('#### Top 10 Highest Risk Suburbs')
        if 'Total_Risk_Score' in valid.columns:
            top_risk = valid.nlargest(10, 'Total_Risk_Score')[[
                'Suburb', 'Current_Price_2025', 'Total_Risk_Score',
            ]].copy()
            if 'Total_Risk_Category' in valid.columns:
                cats = df.set_index('Suburb')['Total_Risk_Category']
                top_risk['Risk Category'] = top_risk['Suburb'].map(cats)
            top_risk['Current_Price_2025'] = top_risk['Current_Price_2025'].apply(fmt_price)
            top_risk['Total_Risk_Score'] = top_risk['Total_Risk_Score'].apply(
                lambda x: fmt_num(x, 1))
            top_risk.columns = ['Suburb', 'Price 2025', 'Risk Score'] + (
                ['Risk Category'] if 'Total_Risk_Category' in valid.columns else [])
            st.dataframe(top_risk, use_container_width=True, hide_index=True)

    st.markdown('---')

    # ── MAP ──
    st.markdown("### Adelaide Property Map")
    st.caption("Click a suburb on the map to see its full report below. Hover for quick stats.")

    if geojson_data or coords:
        m = create_map(df, coords, geojson_data)
        st_data = st_folium(m, width=None, height=550)

        # Handle map click
        if st_data and st_data.get('last_clicked'):
            click = st_data['last_clicked']
            click_id = f"{click['lat']:.4f}_{click['lng']:.4f}"
            if click_id != st.session_state.last_click_id:
                min_dist = float('inf')
                nearest = None
                for sub, coord in coords.items():
                    d = (coord['lat'] - click['lat'])**2 + (coord['lng'] - click['lng'])**2
                    if d < min_dist:
                        min_dist = d
                        nearest = sub
                if nearest and nearest in suburbs_list:
                    st.session_state.last_click_id = click_id
                    st.session_state.map_suburb = nearest
                    st.rerun()
    else:
        st.info("Map data (GeoJSON/coordinates) not available. "
                "The map will appear once suburb_coordinates.json or adelaide_suburbs.geojson is generated.")

    # ── SUBURB REPORT ──
    if selected_suburb:
        row_data = df[df['Suburb'] == selected_suburb]
        if not row_data.empty:
            row = row_data.iloc[0]
            st.markdown("---")
            render_suburb_report(row, ts)
        else:
            st.error(f"No data found for suburb: {selected_suburb}")


if __name__ == "__main__":
    main()
